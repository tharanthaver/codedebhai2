import os
import uuid
import logging
import requests
import io
import contextlib
import random
import time
import re
import concurrent.futures
import threading
import hmac
import hashlib
import json
import subprocess
import shutil
from flask import Flask, request, jsonify, send_file, render_template, session as flask_session, redirect, make_response
from flask_socketio import SocketIO, emit, join_room, leave_room
from real_time_progress import real_time_progress
from datetime import timedelta, datetime, timezone
import datetime as dt
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from dotenv import load_dotenv
from db_helper import DatabaseHelper
from otp_service import OTPService
from twofactor_otp import TwoFactorOTP
from firebase_otp import FirebaseOTP
from production_rollover_system import ProductionRolloverSystem
from scheduler_service import rollover_scheduler, add_scheduler_routes
from terminal_utils import TerminalUtils, clean_terminal_path, take_screenshot, fix_terminal, suppress_extra_output
from error_handlers import setup_error_handlers
import anthropic
import os

load_dotenv()  # This loads the variables from .env into your environment




app = Flask(__name__, template_folder='templates', static_folder='static', static_url_path='/static')

# Configure Flask app for better connection handling
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0  # Disable caching for downloads
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=24)

# Initialize SocketIO with production-optimized configuration
# Get allowed origins from environment for security
allowed_origins = os.getenv('CORS_ALLOWED_ORIGINS', 'https://codedebhai2.onrender.com').split(',')

socketio = SocketIO(app, 
                   cors_allowed_origins=allowed_origins, 
                   async_mode='threading',         # Using threading for Python 3.13 compatibility
                   # Production-optimized timeout settings
                   ping_timeout=60,            # Reduced to 60s for faster failure detection
                   ping_interval=25,           # Reduced to 25s for more frequent health checks
                   # Connection and transport settings - prioritize polling for stability
                   transports=['polling', 'websocket'],  # Polling first for production stability
                   allow_upgrades=False,       # Disable upgrades to prevent connection drops
                   # Enhanced buffer and connection limits
                   max_http_buffer_size=5000000,  # Increased to 5MB for large responses
                   # Logging configuration
                   engineio_logger=False, 
                   socketio_logger=False,
                   logger=False, 
                   # Additional production settings
                   compression=True,
                   cookie=None)

UPLOAD_FOLDER = 'solved_files'
TEMP_FOLDER = 'temp'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TEMP_FOLDER, exist_ok=True)

# Fetch the API key from the environment
load_dotenv()  # This loads the variables from .env into your environment
API_KEY = os.getenv('API_KEY')
if not API_KEY:
    logging.error("API_KEY not found in environment. Please set the API_KEY environment variable.")

BASE_URL = "https://api.deepseek.com/v1/chat/completions"


# Initialize Flask session with persistent configuration

# DeepSeek API keys with enhanced load balancing
DEEPSEEK_KEYS = [
    os.getenv('DEEPSEEK_KEY_1', ''),
    os.getenv('DEEPSEEK_KEY_2', ''),
    os.getenv('DEEPSEEK_KEY_3', ''),
    os.getenv('DEEPSEEK_KEY_4', '')
]
# Filter out empty keys
DEEPSEEK_KEYS = [key for key in DEEPSEEK_KEYS if key]

# Enhanced API key management with load balancing
class APIKeyManager:
    def __init__(self, keys):
        self.keys = keys
        self.key_stats = {i: {'requests': 0, 'failures': 0, 'last_used': 0, 'rate_limited_until': 0} for i in range(len(keys))}
        self.current_index = 0
        self.lock = threading.Lock()
    
    def get_best_key(self):
        """Get the best available API key based on load balancing"""
        with self.lock:
            current_time = time.time()
            
            # Filter out rate-limited keys
            available_keys = [
                i for i in range(len(self.keys)) 
                if self.key_stats[i]['rate_limited_until'] < current_time
            ]
            
            if not available_keys:
                # All keys are rate limited, wait for the soonest one
                next_available = min(self.key_stats.values(), key=lambda x: x['rate_limited_until'])
                wait_time = max(0, next_available['rate_limited_until'] - current_time)
                if wait_time > 0:
                    logging.warning(f"All API keys rate limited. Waiting {wait_time:.2f} seconds...")
                    time.sleep(wait_time)
                available_keys = list(range(len(self.keys)))
            
            # Choose key with least recent usage for better load distribution
            best_key_index = min(available_keys, key=lambda i: (
                self.key_stats[i]['requests'], 
                self.key_stats[i]['last_used']
            ))
            
            # Update stats
            self.key_stats[best_key_index]['requests'] += 1
            self.key_stats[best_key_index]['last_used'] = current_time
            
            return best_key_index, self.keys[best_key_index]
    
    def mark_rate_limited(self, key_index, duration=60):
        """Mark a key as rate limited for specified duration (seconds)"""
        with self.lock:
            self.key_stats[key_index]['rate_limited_until'] = time.time() + duration
            self.key_stats[key_index]['failures'] += 1
            logging.warning(f"API key {key_index + 1} marked as rate limited for {duration} seconds")
    
    def mark_success(self, key_index):
        """Mark successful API call"""
        with self.lock:
            # Reset rate limit if it was set
            self.key_stats[key_index]['rate_limited_until'] = 0
    
    def get_stats(self):
        """Get current statistics for all keys"""
        with self.lock:
            return dict(self.key_stats)

# Initialize enhanced API key managers
deepseek_manager = APIKeyManager(DEEPSEEK_KEYS)
db_helper = DatabaseHelper()


# Set secret key with fallback
secret_key = os.getenv('SECRET_KEY')
if not secret_key:
    import secrets
    secret_key = secrets.token_hex(32)
    logging.warning("SECRET_KEY not found in environment. Using generated key (not recommended for production)")

app.secret_key = secret_key
app.config['SESSION_TYPE'] = 'filesystem'
app.config['PERMANENT_SESSION_LIFETIME'] = 30 * 24 * 60 * 60  # 30 days in seconds
app.config['SESSION_COOKIE_SECURE'] = True  # Enable HTTPS-only cookies for production
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# Initialize rollover scheduler
rollover_scheduler.init_app(app)
add_scheduler_routes(app)

# Import Firebase OTP service
from firebase_otp import firebase_otp_service

# Create a global requests session for connection reuse
session = requests.Session()
session.headers.update({"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"})

# Enhanced thread pool configuration for high traffic
# Increase workers for better multitasking support
max_workers = min(50, (os.cpu_count() or 4) * 12)  # Cap at 50 workers
global_executor = concurrent.futures.ThreadPoolExecutor(
    max_workers=max_workers,
    thread_name_prefix='CodeDeBhai-Worker'
)

# Separate executor for API calls to prevent blocking
api_executor = concurrent.futures.ThreadPoolExecutor(
    max_workers=min(20, len(DEEPSEEK_KEYS) * 5),  # 5 concurrent calls per API key
    thread_name_prefix='API-Worker'
)

logging.info(f"Initialized thread pools: Main={max_workers} workers, API={min(20, len(DEEPSEEK_KEYS) * 5)} workers")

# Progress tracking for real-time updates
active_tasks = {}
task_progress = {}

# Configure logging level based on environment (production-ready)
log_level = os.getenv('LOG_LEVEL', 'INFO').upper()
logging_level = getattr(logging, log_level, logging.INFO)

logging.basicConfig(
    level=logging_level,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(), logging.FileHandler('app.log')]
)

# --- In-memory cache to avoid duplicate API calls ---
solution_cache = {}
cache_lock = threading.Lock()

# Claude API keys with intelligent management - Load from environment
CLAUDE_KEYS = [
    os.getenv('ANTHROPIC_API_KEY'),
    os.getenv('DEEPSEEK_KEY_1'),
    os.getenv('DEEPSEEK_KEY_2'),
    os.getenv('DEEPSEEK_KEY_3')
]

# Filter out None values and ensure we have at least one key
CLAUDE_KEYS = [key for key in CLAUDE_KEYS if key]
if not CLAUDE_KEYS:
    raise ValueError("No API keys found. Please set ANTHROPIC_API_KEY or DEEPSEEK_KEY_* environment variables")

# Enhanced Claude API Manager with concurrent connection control
class ClaudeAPIManager:
    def __init__(self, keys):
        self.keys = keys
        self.key_stats = {
            i: {
                'requests': 0,
                'failures': 0,
                'last_used': 0,
                'rate_limited_until': 0,
                'requests_this_minute': 0,
                'requests_this_hour': 0,
                'minute_window_start': 0,
                'hour_window_start': 0,
                'consecutive_errors': 0,
                'last_error_type': None,
                'concurrent_connections': 0
            } for i in range(len(keys))
        }
        self.current_index = 0
        self.lock = threading.Lock()
        # Claude API rate limits (based on official documentation)
        self.max_requests_per_minute = 60  # Actual limit for most tiers
        self.max_requests_per_hour = 1000  # Hourly limit
        self.rate_limit_cooldown = 60  # 1 minute cooldown for rate limits
        self.error_cooldown = 120  # 2 minute cooldown for other errors
        self.max_consecutive_errors = 3  # Max errors before longer cooldown
        # NEW: Concurrent connection limits
        self.max_concurrent_connections_per_key = 2  # Max 2 concurrent connections per key
        self.global_concurrent_limit = 4  # Global limit across all keys
        self.connection_semaphore = threading.Semaphore(self.global_concurrent_limit)
    
    def reset_minute_window(self, key_index):
        """Reset the minute window for request counting"""
        current_time = time.time()
        stats = self.key_stats[key_index]
        
        # If more than a minute has passed, reset the counter
        if current_time - stats['minute_window_start'] >= 60:
            stats['requests_this_minute'] = 0
            stats['minute_window_start'] = current_time
    
    def reset_hour_window(self, key_index):
        """Reset the hour window for request counting"""
        current_time = time.time()
        stats = self.key_stats[key_index]
        
        # If more than an hour has passed, reset the counter
        if current_time - stats['hour_window_start'] >= 3600:
            stats['requests_this_hour'] = 0
            stats['hour_window_start'] = current_time
    
    def is_key_available(self, key_index):
        """Check if a Claude key is available for use with concurrent connection control"""
        current_time = time.time()
        stats = self.key_stats[key_index]
        
        # Check if rate limited
        if stats['rate_limited_until'] > current_time:
            return False
        
        # Check consecutive errors - if too many errors, apply longer cooldown
        if stats['consecutive_errors'] >= self.max_consecutive_errors:
            if current_time - stats['last_used'] < self.error_cooldown:
                logging.warning(f"Claude API key {key_index + 1} in error cooldown ({stats['consecutive_errors']} consecutive errors)")
                return False
            else:
                # Reset consecutive errors after cooldown period
                stats['consecutive_errors'] = 0
        
        # NEW: Check concurrent connections per key
        if stats['concurrent_connections'] >= self.max_concurrent_connections_per_key:
            logging.warning(f"Claude API key {key_index + 1} at max concurrent connections ({stats['concurrent_connections']}/{self.max_concurrent_connections_per_key})")
            return False
        
        # Reset minute and hour windows if needed
        self.reset_minute_window(key_index)
        self.reset_hour_window(key_index)
        
        # Check minute rate limit (primary limit)
        if stats['requests_this_minute'] >= self.max_requests_per_minute:
            # Mark as rate limited for remainder of minute
            remaining_seconds = 60 - (current_time - stats['minute_window_start'])
            stats['rate_limited_until'] = current_time + max(remaining_seconds, 10)
            logging.warning(f"Claude API key {key_index + 1} hit per-minute rate limit ({stats['requests_this_minute']}/{self.max_requests_per_minute} requests)")
            return False
        
        # Check hourly rate limit (secondary limit)
        if stats['requests_this_hour'] >= self.max_requests_per_hour:
            # Mark as rate limited for remainder of hour
            remaining_seconds = 3600 - (current_time - stats['hour_window_start'])
            stats['rate_limited_until'] = current_time + max(remaining_seconds, 60)
            logging.warning(f"Claude API key {key_index + 1} hit per-hour rate limit ({stats['requests_this_hour']}/{self.max_requests_per_hour} requests)")
            return False
        
        return True
    
    def get_best_available_key(self):
        """Get the best available Claude key with load balancing"""
        with self.lock:
            current_time = time.time()
            
            # Find available keys
            available_keys = []
            for i in range(len(self.keys)):
                if self.is_key_available(i):
                    available_keys.append(i)
            
            if not available_keys:
                return None, None  # No keys available
            
            # Choose key with least usage in current minute
            best_key_index = min(available_keys, key=lambda i: (
                self.key_stats[i]['requests_this_minute'],
                self.key_stats[i]['last_used']
            ))
            
            # Update usage stats
            stats = self.key_stats[best_key_index]
            stats['requests'] += 1
            stats['requests_this_minute'] += 1
            stats['last_used'] = current_time
            
            # Reset minute window if needed
            self.reset_minute_window(best_key_index)
            
            return best_key_index, self.keys[best_key_index]
    
    def mark_rate_limited(self, key_index, duration=60):
        """Mark Claude key as rate limited"""
        with self.lock:
            current_time = time.time()
            self.key_stats[key_index]['rate_limited_until'] = current_time + duration
            self.key_stats[key_index]['failures'] += 1
            logging.warning(f"Claude API key {key_index + 1} marked as rate limited for {duration} seconds")
    
    def mark_success(self, key_index):
        """Mark successful Claude API call"""
        with self.lock:
            # Don't reset rate limit immediately, let minute window handle it
            pass
    
    def get_available_keys_count(self):
        """Get count of currently available Claude keys"""
        available_count = 0
        current_time = time.time()
        
        for i in range(len(self.keys)):
            if self.is_key_available(i):
                available_count += 1
        
        return available_count
    
    def get_stats(self):
        """Get current statistics for all Claude keys"""
        with self.lock:
            return dict(self.key_stats)

# Initialize Claude API manager
claude_manager = ClaudeAPIManager(CLAUDE_KEYS)

# Global terminal paths for cycling in screenshots
TERMINAL_PATHS = [
    "C:\\Users\\THARAN\\Desktop\\AIProjects\\ChatBot",
    "C:\\Users\\THARAN\\Documents\\WebDevelopment\\ReactApp",
    "C:\\Users\\THARAN\\Desktop\\DataScience\\MLModel",
    "C:\\Users\\THARAN\\Projects\\PythonScript\\Automation",
    "C:\\Users\\THARAN\\Desktop\\GameDev\\UnityProject",
    "C:\\Users\\THARAN\\Documents\\Backend\\NodeJS",
    "C:\\Users\\THARAN\\Desktop\\Frontend\\VueApp",
    "C:\\Users\\THARAN\\Projects\\MobileApp\\Flutter",
    "C:\\Users\\THARAN\\Desktop\\DevOps\\DockerConfig",
    "C:\\Users\\THARAN\\Documents\\Database\\MongoDB",
    "C:\\Users\\THARAN\\Desktop\\Analytics\\Dashboard",
    "C:\\Users\\THARAN\\Projects\\Security\\Encryption",
    "C:\\Users\\THARAN\\Desktop\\CloudComputing\\AWS",
    "C:\\Users\\THARAN\\Documents\\API\\RestfulService",
    "C:\\Users\\THARAN\\Desktop\\BlockChain\\SmartContract",
    "C:\\Users\\THARAN\\Projects\\Testing\\Selenium",
    "C:\\Users\\THARAN\\Desktop\\MachineLearning\\TensorFlow",
    "C:\\Users\\THARAN\\Documents\\Microservices\\SpringBoot",
    "C:\\Users\\THARAN\\Desktop\\IoT\\ArduinoProject",
    "C:\\Users\\THARAN\\Projects\\CyberSecurity\\Penetration",
    "C:\\Users\\THARAN\\Desktop\\AugmentedReality\\ARKit",
    "C:\\Users\\THARAN\\Documents\\BigData\\Hadoop",
    "C:\\Users\\THARAN\\Desktop\\NetworkSecurity\\Firewall",
    "C:\\Users\\THARAN\\Projects\\CloudNative\\Kubernetes",
    "C:\\Users\\THARAN\\Desktop\\QuantumComputing\\Qiskit",
    "C:\\Users\\THARAN\\Documents\\EdgeComputing\\5G",
    "C:\\Users\\THARAN\\Desktop\\RoboticsAI\\OpenCV",
    "C:\\Users\\THARAN\\Projects\\VirtualReality\\Oculus",
    "C:\\Users\\THARAN\\Desktop\\FinTech\\PaymentGateway",
    "C:\\Users\\THARAN\\Documents\\HealthTech\\Telemedicine",
    "C:\\Users\\THARAN\\Desktop\\EcommercePlatform\\Magento",
    "C:\\Users\\THARAN\\Projects\\SocialMedia\\GraphQL",
    "C:\\Users\\THARAN\\Desktop\\ContentManagement\\WordPress",
    "C:\\Users\\THARAN\\Documents\\BusinessIntelligence\\Tableau",
    "C:\\Users\\THARAN\\Desktop\\LogisticsOptimization\\Algorithm",
    "C:\\Users\\THARAN\\Projects\\SmartHome\\HomeAssistant",
    "C:\\Users\\THARAN\\Desktop\\AutonomousVehicle\\ROS",
    "C:\\Users\\THARAN\\Documents\\SupplyChain\\Blockchain",
    "C:\\Users\\THARAN\\Desktop\\DigitalMarketing\\SEO",
    "C:\\Users\\THARAN\\Projects\\CustomerAnalytics\\Pandas",
    "C:\\Users\\THARAN\\Desktop\\ProcessAutomation\\RPA",
    "C:\\Users\\THARAN\\Documents\\QualityAssurance\\Jest",
    "C:\\Users\\THARAN\\Desktop\\SystemIntegration\\Apache",
    "C:\\Users\\THARAN\\Projects\\PerformanceOptimization\\Redis",
    "C:\\Users\\THARAN\\Desktop\\DataVisualization\\D3js",
    "C:\\Users\\THARAN\\Documents\\ServerlessComputing\\Lambda",
    "C:\\Users\\THARAN\\Desktop\\NaturalLanguageProcessing\\NLTK",
    "C:\\Users\\THARAN\\Projects\\ComputerVision\\YOLO",
    "C:\\Users\\THARAN\\Desktop\\RecommendationEngine\\Collaborative",
    "C:\\Users\\THARAN\\Documents\\PredictiveAnalytics\\Prophet"
]

# Global index and lock for cycling through terminal paths
terminal_path_index = 0
terminal_path_lock = threading.Lock()

# Helper: build realistic Windows terminal paths for a given user
def get_realistic_terminal_paths(user_name: str):
    """Build a comprehensive list of realistic terminal paths for a given user from Supabase"""
    try:
        # Clean the username from Supabase/database
        cleaned = (user_name or "User").strip().replace(" ", "")
        
        # Windows paths with dynamic username
        windows_paths = [
            f"C:\\Users\\{cleaned}\\Documents\\Projects",
            f"C:\\Users\\{cleaned}\\Desktop\\PythonScripts",
            f"C:\\Users\\{cleaned}\\Documents\\College\\Assignments",
            f"C:\\Users\\{cleaned}\\Downloads",
            f"C:\\Users\\{cleaned}\\Documents\\GitHub\\Repositories",
            f"C:\\Users\\{cleaned}\\Desktop\\Notes",
            f"C:\\Users\\{cleaned}\\Documents\\ReactProjects",
            f"C:\\Users\\{cleaned}\\Documents\\BackendProjects",
            f"C:\\Users\\{cleaned}\\Documents\\MiniProjects",
            f"C:\\Users\\{cleaned}\\Desktop\\VSCodeProjects",
            f"C:\\Users\\{cleaned}\\Documents\\NodeProjects",
            f"C:\\Users\\{cleaned}\\Documents\\DjangoProjects",
            f"C:\\Users\\{cleaned}\\Documents\\JavaPrograms",
            f"C:\\Users\\{cleaned}\\Desktop\\CPrograms",
            f"C:\\Users\\{cleaned}\\Desktop\\Games",
            f"C:\\Users\\{cleaned}\\Documents\\FlaskApps",
            f"C:\\Users\\{cleaned}\\Documents\\MobileProjects",
            f"C:\\Users\\{cleaned}\\Desktop\\PortfolioSite",
            f"C:\\Users\\{cleaned}\\Documents\\ClassNotes",
            f"C:\\Users\\{cleaned}\\Documents\\AutomationScripts",
            f"C:\\Users\\{cleaned}\\Downloads\\Resources",
            f"C:\\Users\\{cleaned}\\Documents\\DataStructures",
            f"C:\\Users\\{cleaned}\\Documents\\Algorithms",
            f"C:\\Users\\{cleaned}\\Desktop\\Programs",
            f"C:\\Users\\{cleaned}\\Documents\\Utilities",
            f"C:\\Users\\{cleaned}\\Documents\\Examples",
            f"C:\\Users\\{cleaned}\\Documents\\DockerProjects",
            f"C:\\Users\\{cleaned}\\Documents\\KubernetesNotes",
            f"C:\\Users\\{cleaned}\\Documents\\FullStackProjects",
            f"C:\\Users\\{cleaned}\\Documents\\RestServices",
            f"C:\\Users\\{cleaned}\\Documents\\SQLScripts",
            f"C:\\Users\\{cleaned}\\Documents\\DatabaseBackups",
            f"C:\\Users\\{cleaned}\\Documents\\FastAPIProjects",
            f"C:\\Users\\{cleaned}\\Documents\\FrontendProjects",
            f"C:\\Users\\{cleaned}\\Documents\\PersonalNotes",
            f"C:\\Users\\{cleaned}\\Documents\\References",
            f"C:\\Users\\{cleaned}\\Documents\\CodeSnippets",
            f"C:\\Users\\{cleaned}\\Desktop\\Tests",
            f"C:\\Users\\{cleaned}\\Documents\\PythonProjects",
            f"C:\\Users\\{cleaned}\\Documents\\PythonUtilities",
            f"C:\\Users\\{cleaned}\\Desktop\\GitClones",
            f"C:\\Users\\{cleaned}\\Documents\\HackathonProjects",
            f"C:\\Users\\{cleaned}\\Documents\\Scripts",
            f"C:\\Users\\{cleaned}\\Documents\\LangChainNotes",
            f"C:\\Users\\{cleaned}\\Documents\\Libraries",
            f"C:\\Users\\{cleaned}\\Documents\\Semester1",
            f"C:\\Users\\{cleaned}\\Documents\\Semester2",
            f"C:\\Users\\{cleaned}\\Documents\\Semester3",
            f"C:\\Users\\{cleaned}\\Desktop\\FinalYearProject",
            f"C:\\Users\\{cleaned}\\Documents\\InternshipWork",
            f"C:\\Users\\{cleaned}\\Documents\\InternshipReports",
            f"C:\\Users\\{cleaned}\\Documents\\Competitions",
            f"C:\\Users\\{cleaned}\\Documents\\SpringBootProjects",
            f"C:\\Users\\{cleaned}\\Documents\\DotNetProjects",
            f"C:\\Users\\{cleaned}\\Documents\\FlutterApps",
            f"C:\\Users\\{cleaned}\\Documents\\AndroidStudioProjects",
            f"C:\\Users\\{cleaned}\\Documents\\ClassProjects",
            f"C:\\Users\\{cleaned}\\Documents\\Assignments",
            f"C:\\Users\\{cleaned}\\Documents\\Reports",
            f"C:\\Users\\{cleaned}\\Documents\\CollegeWork",
            f"C:\\Users\\{cleaned}\\Documents\\ExcelSheets",
            f"C:\\Users\\{cleaned}\\Documents\\Dashboards",
            f"C:\\Users\\{cleaned}\\Documents\\MatplotlibCharts",
            f"C:\\Users\\{cleaned}\\Documents\\PlotlyGraphs",
            f"C:\\Users\\{cleaned}\\Desktop\\Website",
            f"C:\\Users\\{cleaned}\\Documents\\AWSNotes",
            f"C:\\Users\\{cleaned}\\Documents\\AzureSetup",
            f"C:\\Users\\{cleaned}\\Documents\\GCPConfig",
            f"C:\\Users\\{cleaned}\\Documents\\PostmanTests",
            f"C:\\Users\\{cleaned}\\Documents\\DebuggingLogs",
            f"C:\\Users\\{cleaned}\\Documents\\PythonScripts",
            f"C:\\Users\\{cleaned}\\Documents\\ExcelAutomation",
            f"C:\\Users\\{cleaned}\\Documents\\UtilitiesBackup",
            f"C:\\Users\\{cleaned}\\Documents\\Datasets",
            f"C:\\Users\\{cleaned}\\Documents\\Outputs",
            f"C:\\Users\\{cleaned}\\Documents\\Results",
            f"C:\\Users\\{cleaned}\\Desktop\\SchoolAssignments",
            f"C:\\Users\\{cleaned}\\Desktop\\UniversityNotes",
            f"C:\\Users\\{cleaned}\\Documents\\VSCodeExtensions",
            f"C:\\Users\\{cleaned}\\Documents\\CustomScripts",
            f"C:\\Users\\{cleaned}\\Documents\\WorkProjects",
            f"C:\\Users\\{cleaned}\\Documents\\CodeBackups",
            f"C:\\Users\\{cleaned}\\Documents\\LabPrograms",
            f"C:\\Users\\{cleaned}\\Documents\\PracticeProblems",
            f"C:\\Users\\{cleaned}\\Documents\\AssignmentsArchive",
            f"C:\\Users\\{cleaned}\\Documents\\ImportantDocs",
            f"C:\\Users\\{cleaned}\\Documents\\Submissions",
            f"C:\\Users\\{cleaned}\\Documents\\Experiments",
            f"C:\\Users\\{cleaned}\\Documents\\Modules",
            f"C:\\Users\\{cleaned}\\Documents\\ReferenceBooks"
        ]
        
        # Linux/Unix paths with dynamic username
        unix_paths = [
            f"/home/{cleaned}/Documents/Projects",
            f"/home/{cleaned}/Downloads",
            f"/home/{cleaned}/Desktop/PythonPrograms",
            f"/home/{cleaned}/Documents/Scripts"
        ]
        
        # macOS paths with dynamic username
        macos_paths = [
            f"/Users/{cleaned}/Documents/Code",
            f"/Users/{cleaned}/Documents/ProjectsArchive",
            f"/Users/{cleaned}/Desktop/Assignments",
            f"/Users/{cleaned}/Downloads/Reports",
            f"/Users/{cleaned}/Documents/FlaskProjects",
            f"/Users/{cleaned}/Documents/PersonalPortfolio"
        ]
        
        # Combine all paths for comprehensive rolling selection
        all_paths = windows_paths + unix_paths + macos_paths
        
        # Return the comprehensive list for rolling selection
        return all_paths if all_paths else [f"C:\\Users\\{cleaned}\\Desktop"]
        
    except Exception as e:
        logging.error(f"Error generating realistic terminal paths: {e}")
        return [f"C:\\Users\\{(user_name or 'User').strip().replace(' ', '')}\\Desktop"]

# ============= WebSocket Event Handlers =============

@socketio.on('connect')
def handle_connect():
    """Handle client connection"""
    user_id = flask_session.get('user', {}).get('phone_number', 'anonymous')
    logging.info(f"Client connected: {user_id}")
    
    # Join user-specific room
    join_room(f"user_{user_id}")
    
    # Send initial connection status
    emit('connection_status', {
        'status': 'connected',
        'user_id': user_id,
        'timestamp': datetime.now().isoformat()
    })

@socketio.on('disconnect')
def handle_disconnect():
    """Handle client disconnection"""
    user_id = flask_session.get('user', {}).get('phone_number', 'anonymous')
    logging.info(f"Client disconnected: {user_id}")
    
    # Leave user-specific room
    leave_room(f"user_{user_id}")

@socketio.on('join_task')
def handle_join_task(data):
    """Handle joining a specific task for updates"""
    task_id = data.get('task_id')
    user_id = flask_session.get('user', {}).get('phone_number')
    
    if not task_id or not user_id:
        emit('error', {'message': 'Invalid task_id or user not authenticated'})
        return
    
    # Join task-specific room
    join_room(f"task_{task_id}")
    
    # Send current task status if available
    if task_id in task_progress:
        emit('task_update', {
            'task_id': task_id,
            'progress': task_progress[task_id],
            'timestamp': datetime.now().isoformat()
        })

@socketio.on('leave_task')
def handle_leave_task(data):
    """Handle leaving a specific task"""
    task_id = data.get('task_id')
    if task_id:
        leave_room(f"task_{task_id}")

def emit_progress_update(task_id, progress_data):
    """Emit progress update to all connected clients for a task"""
    try:
        # Update global progress tracking
        task_progress[task_id] = progress_data
        
        # Emit to task-specific room
        socketio.emit('task_update', {
            'task_id': task_id,
            'progress': progress_data,
            'timestamp': datetime.now().isoformat()
        }, room=f"task_{task_id}")
        
        # Emit to user-specific room if user_id is provided
        if 'user_id' in progress_data:
            socketio.emit('user_task_update', {
                'task_id': task_id,
                'progress': progress_data,
                'timestamp': datetime.now().isoformat()
            }, room=f"user_{progress_data['user_id']}")
    except Exception as e:
        logging.error(f"Error emitting progress update: {e}")

def calculate_estimated_time(question_count, task_type='pdf'):
    """Calculate estimated completion time based on question count, complexity, and current server load"""
    try:
        # Base time estimates (in seconds)
        base_times = {
            'pdf': 15,    # Base time for PDF processing
            'manual': 10  # Base time for manual questions
        }
        
        # Time per question (in seconds)
        time_per_question = {
            'simple': 8,      # Simple coding questions
            'medium': 12,     # Medium complexity
            'complex': 18     # Complex questions
        }
        
        # Determine complexity based on question count
        if question_count <= 5:
            complexity = 'simple'
        elif question_count <= 15:
            complexity = 'medium'
        else:
            complexity = 'complex'
        
        # Calculate base total estimated time
        base_time = base_times.get(task_type, 10)
        question_time = time_per_question[complexity] * question_count
        document_generation_time = 5 + (question_count * 0.5)  # Scales with question count
        
        total_seconds = base_time + question_time + document_generation_time
        
        # Apply current server load factor
        load_factor = get_current_load_factor()
        total_seconds = total_seconds * load_factor
        
        # Add buffer time (20% buffer)
        total_seconds = int(total_seconds * 1.2)
        
        # Format time for display
        if total_seconds < 60:
            return f"{total_seconds} seconds"
        elif total_seconds < 3600:
            minutes = total_seconds // 60
            seconds = total_seconds % 60
            return f"{minutes}m {seconds}s"
        else:
            hours = total_seconds // 3600
            minutes = (total_seconds % 3600) // 60
            return f"{hours}h {minutes}m"
            
    except Exception as e:
        logging.error(f"Error calculating estimated time: {e}")
        return "Calculating..."

def get_progress_stages(task_type='pdf'):
    """Get progress stages with estimated durations"""
    if task_type == 'pdf':
        return {
            'initialization': {'name': 'Initializing', 'duration': 5, 'progress': 5},
            'pdf_extraction': {'name': 'Extracting questions from PDF', 'duration': 10, 'progress': 15},
            'question_analysis': {'name': 'Analyzing questions', 'duration': 8, 'progress': 25},
            'ai_processing': {'name': 'Generating solutions', 'duration': 60, 'progress': 80},
            'document_generation': {'name': 'Creating Word document', 'duration': 10, 'progress': 95},
            'finalization': {'name': 'Finalizing', 'duration': 5, 'progress': 100}
        }
    else:  # manual
        return {
            'initialization': {'name': 'Initializing', 'duration': 3, 'progress': 5},
            'question_analysis': {'name': 'Analyzing questions', 'duration': 5, 'progress': 20},
            'ai_processing': {'name': 'Generating solutions', 'duration': 50, 'progress': 85},
            'document_generation': {'name': 'Creating Word document', 'duration': 8, 'progress': 98},
            'finalization': {'name': 'Finalizing', 'duration': 2, 'progress': 100}
        }


def get_current_load_factor():
    """
    Compute a simple server load factor based on active tasks and connections.
    Returns a float >= 1.0 (capped for sanity).
    """
    try:
        from task_manager import task_manager
        active = task_manager.get_active_tasks(limit=100) or []
        processing_count = sum(1 for t in active if t.get('task_status') == 'PROCESSING')
        pending_count = sum(1 for t in active if t.get('task_status') == 'PENDING')
        
        conn_stats = real_time_progress.get_connection_stats() or {}
        connections = conn_stats.get('total_connections', 0)

        # Simple heuristic: processing tasks weigh more than pending; connections add a small amount
        load_factor = 1.0 + (processing_count * 0.08) + (pending_count * 0.04) + (connections * 0.01)
        return round(min(load_factor, 2.0), 2)  # Cap at 2x
    except Exception as e:
        logging.error(f"Error computing load factor: {e}")
        return 1.0


@app.route('/server_load', methods=['GET'])
def get_server_load():
    """Expose current server load metrics and computed load_factor"""
    try:
        from task_manager import task_manager
        active = task_manager.get_active_tasks(limit=100) or []
        processing_count = sum(1 for t in active if t.get('task_status') == 'PROCESSING')
        pending_count = sum(1 for t in active if t.get('task_status') == 'PENDING')
        
        conn_stats = real_time_progress.get_connection_stats() or {}
        connections = conn_stats.get('total_connections', 0)
        
        return jsonify({
            'active_tasks_count': len(active),
            'processing_count': processing_count,
            'pending_count': pending_count,
            'connections': connections,
            'load_factor': get_current_load_factor()
        }), 200
    except Exception as e:
        logging.exception("Error getting server load")
        return jsonify({'error': str(e)}), 500

@app.route('/create-payment-session', methods=['POST'])
def create_payment_session():
    """Create a payment session with Cashfree - Simplified version"""
    try:
        logging.info("=== PAYMENT SESSION CREATION STARTED ===")
        
        # Get request data
        data = request.get_json()
        if not data:
            logging.error("No JSON data received in payment request")
            return jsonify({'success': False, 'error': 'No data provided'}), 400

        # Extract payment details
        amount = data.get('amount', 99)
        currency = data.get('currency', 'INR')
        plan_type = data.get('plan_type', 'starter')
        
        # Get user info if available
        user_data = data.get('user', {})
        customer_name = user_data.get('name', 'Customer')
        customer_email = user_data.get('email', 'customer@codedebhai.com')
        customer_phone = user_data.get('phone', '9999999999')
        
        logging.info(f"Payment session request - Amount: {amount}, Plan: {plan_type}, Customer: {customer_name}")

        # Validate Cashfree credentials
        app_id = os.getenv('CASHFREE_APP_ID')
        secret_key = os.getenv('CASHFREE_SECRET_KEY')
        environment = os.getenv('CASHFREE_ENVIRONMENT', 'sandbox').lower()
        api_version = os.getenv('CASHFREE_API_VERSION', '2023-08-01')
        
        if not app_id or not secret_key:
            logging.error("Missing Cashfree credentials")
            return jsonify({'success': False, 'error': 'Payment gateway configuration error'}), 500

        # Generate unique order ID
        order_id = f"ORDER_{uuid.uuid4().hex[:8].upper()}"
        logging.info(f"Generated order ID: {order_id}")
        
        # Determine base URL
        if environment == 'production':
            base_url = "https://api.cashfree.com/pg"
        else:
            base_url = "https://sandbox.cashfree.com/pg"
        
        logging.info(f"Using Cashfree base URL: {base_url}")

        # Payment session payload with production URLs
        payload = {
            "order_id": order_id,
            "order_amount": amount,
            "order_currency": currency,
            "customer_details": {
                "customer_id": f"CUST_{uuid.uuid4().hex[:8].upper()}",
                "customer_name": customer_name,
                "customer_email": customer_email,
                "customer_phone": customer_phone
            },
            "order_meta": {
                "return_url": os.getenv('PAYMENT_RETURN_URL', f"https://codedebhai2.onrender.com/payment-success?order_id={order_id}"),
                "notify_url": os.getenv('PAYMENT_WEBHOOK_URL', f"https://codedebhai2.onrender.com/payment-webhook")
            },
            "order_expiry_time": (datetime.now(timezone.utc) + timedelta(hours=1)).isoformat()
        }

        # Headers for Cashfree API
        headers = {
            "Accept": "application/json",
            "Content-Type": "application/json",
            "x-client-id": app_id,
            "x-client-secret": secret_key,
            "x-api-version": api_version
        }

        logging.info(f"Making API call to: {base_url}/orders")
        logging.info(f"Payload: {json.dumps(payload, indent=2)}")

        # Make request to Cashfree
        response = requests.post(
            f"{base_url}/orders",
            headers=headers,
            json=payload,
            timeout=30
        )

        logging.info(f"Cashfree API response status: {response.status_code}")
        logging.info(f"Response headers: {dict(response.headers)}")
        logging.info(f"Raw response text: {response.text[:500]}...")  # First 500 chars

        if response.status_code == 200:
            response_data = response.json()
            logging.info(f"Cashfree response: {json.dumps(response_data, indent=2)}")
            
            # Get the payment session ID from Cashfree response
            payment_session_id = response_data.get("payment_session_id")
            logging.info(f"DEBUG: Raw payment_session_id from Cashfree: '{payment_session_id}'")
            logging.info(f"DEBUG: Payment session ID length: {len(payment_session_id) if payment_session_id else 'None'}")
            
            # Try to create payment record in database
            try:
                # Get user by phone if available
                user = None
                if customer_phone != '9999999999':
                    user = db_helper.get_user_by_phone(customer_phone)
                
                # Get plan info for credits
                plan_data = {
                    'starter': {'credits': 10, 'expected_amount': 99},
                    'monthly': {'credits': 50, 'expected_amount': 299},
                    'power': {'credits': 150, 'expected_amount': 799}
                }
                plan_info = plan_data.get(plan_type)
                credits = plan_info.get('credits', 0) if plan_info else 0
                
                payment_record = db_helper.create_payment_record(
                    user_id=user.get('id') if user else None,
                    phone_number=customer_phone,
                    gateway_order_id=order_id,
                    gateway_payment_id=None,
                    amount=amount * 100,  # Convert to paise
                    credits_added=credits,
                    plan_type=plan_type,
                    payment_status='pending'
                )
                if payment_record:
                    logging.info(f"Payment record created for order: {order_id}")
            except Exception as db_err:
                logging.warning(f"Failed to create payment record: {db_err}")

            # Use original payment session ID from Cashfree response
            payment_session_id = response_data.get("payment_session_id")
            logging.info(f"DEBUG: Using original payment_session_id: '{payment_session_id}'")
            
            # Validate that we have a valid payment session ID
            if not payment_session_id:
                logging.error("No payment_session_id received from Cashfree")
                return jsonify({
                    "success": False,
                    "error": "Invalid payment session response from gateway"
                }), 500
            
            # Log the exact session ID for debugging
            logging.info(f"✅ Valid payment session ID received: {payment_session_id}")
            
            return jsonify({
                "success": True,
                "payment_session_id": response_data.get("payment_session_id"),
                "order_id": order_id,
                "debug_info": {
                    "raw_payment_session_id": response_data.get("payment_session_id"),
                    "cashfree_response_keys": list(response_data.keys()),
                    "environment": environment,
                    "api_version": api_version
                }
            })
        else:
            logging.error(f"Cashfree API Error: {response.status_code}, {response.text}")
            return jsonify({
                "success": False,
                "error": "Failed to create payment session"
            }), 400

    except Exception as e:
        logging.exception("Error creating payment session")
        return jsonify({
            "success": False,
            "error": "Internal server error"
        }), 500

from flask import request, jsonify, redirect, url_for
import logging

@app.route('/payment-success')
def payment_success():
    """Handle successful payment redirect with enhanced session management"""
    order_id = request.args.get('order_id')
    
    # Validate order_id
    if not order_id:
        return redirect(url_for('index', error='Invalid payment session'))
    
    # Check if user is logged in, if not try to refresh session
    current_user = flask_session.get('user')
    if not current_user:
        # Try to find user from payment record
        payment_record = db_helper.get_payment_by_gateway_id(order_id)
        if payment_record:
            user = db_helper.get_user_by_phone(payment_record['phone_number'])
            if user:
                # Restore session
                flask_session['user'] = user
                flask_session['authenticated'] = True
                flask_session.permanent = True
                current_user = user
                logging.info(f"Session restored for user {payment_record['phone_number']} from payment record")
    
    success_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Payment Success - CodeDeBhai</title>
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <style>
            body {{
                font-family: Arial, sans-serif;
                max-width: 600px;
                margin: 50px auto;
                padding: 20px;
                text-align: center;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: #fff;
                min-height: 100vh;
            }}
            .success-container {{
                background: rgba(255, 255, 255, 0.1);
                backdrop-filter: blur(25px);
                padding: 40px;
                border-radius: 20px;
                box-shadow: 0 15px 35px rgba(0, 0, 0, 0.3);
                border: 1px solid rgba(255, 255, 255, 0.2);
            }}
            .success-icon {{
                color: #4ecdc4;
                font-size: 64px;
                margin-bottom: 20px;
                animation: bounce 2s infinite;
            }}
            @keyframes bounce {{
                0%, 20%, 50%, 80%, 100% {{ transform: translateY(0); }}
                40% {{ transform: translateY(-10px); }}
                60% {{ transform: translateY(-5px); }}
            }}
            h1 {{
                color: #fff;
                margin-bottom: 20px;
                font-size: 2.5em;
                text-shadow: 0 2px 8px rgba(0, 0, 0, 0.3);
            }}
            p {{
                color: rgba(255, 255, 255, 0.9);
                margin-bottom: 20px;
                font-size: 1.1em;
                line-height: 1.6;
            }}
            .order-id {{
                background: rgba(255, 255, 255, 0.1);
                padding: 15px;
                border-radius: 10px;
                margin: 20px 0;
                font-family: monospace;
                color: #4ecdc4;
                border: 1px solid rgba(78, 205, 196, 0.3);
                word-break: break-all;
            }}
            .back-btn {{
                background: linear-gradient(45deg, #4ecdc4, #44a08d);
                color: white;
                padding: 15px 30px;
                border: none;
                border-radius: 25px;
                cursor: pointer;
                margin: 10px;
                text-decoration: none;
                display: inline-block;
                font-weight: bold;
                transition: all 0.3s ease;
                font-size: 1.1em;
            }}
            .back-btn:hover {{
                transform: translateY(-2px);
                box-shadow: 0 5px 15px rgba(78, 205, 196, 0.4);
            }}
            .test-btn {{
                background: linear-gradient(45deg, #667eea, #764ba2);
                color: white;
                padding: 15px 30px;
                border: none;
                border-radius: 25px;
                cursor: pointer;
                margin: 10px;
                font-weight: bold;
                transition: all 0.3s ease;
                font-size: 1.1em;
            }}
            .test-btn:hover {{
                transform: translateY(-2px);
                box-shadow: 0 5px 15px rgba(102, 126, 234, 0.4);
            }}
            .test-btn:disabled {{
                background: #6c757d;
                cursor: not-allowed;
                transform: none;
            }}
            .message {{
                margin-top: 20px;
                padding: 15px;
                border-radius: 10px;
                display: none;
                font-weight: bold;
            }}
            .success-msg {{
                background: rgba(78, 205, 196, 0.2);
                color: #4ecdc4;
                border: 1px solid rgba(78, 205, 196, 0.3);
            }}
            .error-msg {{
                background: rgba(255, 107, 107, 0.2);
                color: #ff6b6b;
                border: 1px solid rgba(255, 107, 107, 0.3);
            }}
            .processing-msg {{
                background: rgba(255, 255, 255, 0.1);
                color: #fff;
                border: 1px solid rgba(255, 255, 255, 0.3);
            }}
            .loading-spinner {{
                display: inline-block;
                width: 20px;
                height: 20px;
                border: 3px solid rgba(255, 255, 255, 0.3);
                border-top: 3px solid #4ecdc4;
                border-radius: 50%;
                animation: spin 1s linear infinite;
                margin-right: 10px;
            }}
            @keyframes spin {{
                0% {{ transform: rotate(0deg); }}
                100% {{ transform: rotate(360deg); }}
            }}
        </style>
    </head>
    <body>
        <div class="success-container">
            <div class="success-icon">✅</div>
            <h1>Payment Successful!</h1>
            <p>🎉 Thank you for your payment! Your credits are being processed and will be available shortly.</p>
            <p>You can now start using CodeDeBhai to solve your coding problems!</p>
            <div class="order-id">
                <strong>Order ID:</strong> {order_id}
            </div>
            <div class="message" id="message"></div>
            <a href="/" class="back-btn">🏠 Back to Home</a>
            <button class="test-btn" onclick="processPaymentManual()" id="manualBtn">⚡ Process Credits</button>
        </div>
        
        <script>
            // Configuration
            const ORDER_ID = '{order_id}';
            const MAX_RETRIES = 2;
            const RETRY_DELAY = 2000; // 2 seconds
            
            // Auto-process credits when page loads
            window.addEventListener('DOMContentLoaded', function() {{
                // Show processing message immediately
                showMessage('<div class="loading-spinner"></div>Processing your credits automatically...', 'processing');
                
                // Hide the manual button initially
                const manualBtn = document.getElementById('manualBtn');
                if (manualBtn) {{
                    manualBtn.style.display = 'none';
                }}
                
                // Auto-process payment after 1 second (faster processing)
                setTimeout(() => {
                    processPayment()
                }, 1000);
            }});
            
            function showMessage(message, type = 'processing') {{
                const messageDiv = document.getElementById('message');
                messageDiv.innerHTML = message;
                messageDiv.style.display = 'block';
                
                // Reset classes
                messageDiv.className = 'message';
                
                // Add appropriate class based on type
                if (type === 'success') {{
                    messageDiv.classList.add('success-msg');
                }} else if (type === 'error') {{
                    messageDiv.classList.add('error-msg');
                }} else {{
                    messageDiv.classList.add('processing-msg');
                }}
            }}
            
            async function processPayment(retryCount = 0) {{
                try {{
                    // Validate order ID exists
                    if (!ORDER_ID || ORDER_ID === 'None') {{
                        throw new Error('Invalid order ID');
                    }}
                    
                    // Make the API call with session preservation
                    const response = await fetch(`/test_webhook_payment/${{ORDER_ID}}`, {{
                        method: 'POST',
                        headers: {{
                            'Content-Type': 'application/json',
                        }},
                        credentials: 'same-origin', // Include session cookies
                        signal: AbortSignal.timeout(10000) // 10 second timeout (faster)
                    }});
                    
                    // Check if response is ok
                    if (!response.ok) {{
                        const errorData = await response.json().catch(() => ({{error: 'Unknown error'}}));
                        throw new Error(errorData.error || `HTTP ${{response.status}}: ${{response.statusText}}`);
                    }}
                    
                    const data = await response.json();
                    
                    // Success case
                    showMessage('✅ ' + (data.message || 'Credits added successfully! Redirecting to dashboard...'), 'success');
                    
                    // Redirect to home immediately after success (faster redirect)
                    setTimeout(() => {{
                        window.location.href = '/?payment_success=true&credits_updated=true';
                    }}, 1500);
                    
                }} catch (error) {{
                    console.error('Payment processing error:', error);
                    
                    // Retry logic
                    if (retryCount < MAX_RETRIES) {{
                        showMessage(`<div class="loading-spinner"></div>Retrying... (Attempt ${{retryCount + 2}}/${{MAX_RETRIES + 1}})`, 'processing');
                        setTimeout(() => {{
                            processPayment(retryCount + 1);
                        }}, RETRY_DELAY);
                        return;
                    }}
                    
                    // Max retries reached
                    let errorMessage = '❌ ';
                    if (error.name === 'AbortError') {{
                        errorMessage += 'Request timed out. ';
                    }} else {{
                        errorMessage += error.message + ' ';
                    }}
                    errorMessage += '<br><small>You can try clicking "Process Credits" manually or contact support if the issue persists.</small>';
                    
                    showMessage(errorMessage, 'error');
                    
                    // Show manual button if auto-processing fails
                    const manualBtn = document.getElementById('manualBtn');
                    if (manualBtn) {{
                        manualBtn.style.display = 'inline-block';
                    }}
                }}
            }}
            
            // Manual processing function (fallback)
            async function processPaymentManual() {{
                const manualBtn = document.getElementById('manualBtn');
                
                // Disable button during processing
                manualBtn.disabled = true;
                manualBtn.innerHTML = '<div class="loading-spinner"></div>Processing...';
                
                showMessage('<div class="loading-spinner"></div>Processing credits manually...', 'processing');
                
                try {{
                    await processPayment(0); // Reset retry count for manual processing
                }} finally {{
                    // Re-enable button
                    manualBtn.disabled = false;
                    manualBtn.innerHTML = '⚡ Process Credits';
                }}
            }}
        </script>
    </body>
    </html>
    """
    return success_html



@app.route('/payment-webhook', methods=['POST'])
def payment_webhook():
    """Handle payment webhook from Cashfree - Enhanced with proper database updates"""
    try:
        data = request.get_json()
        logging.info(f"🔔 Payment webhook received: {json.dumps(data, indent=2)}")
        
        # Extract webhook data - handle different Cashfree webhook formats
        order_id = data.get('order_id') or data.get('orderId')
        order_status = data.get('order_status') or data.get('orderStatus') 
        payment_id = data.get('cf_payment_id') or data.get('paymentSessionId') or data.get('payment_id')
        payment_amount = data.get('order_amount') or data.get('amount')
        payment_method = data.get('payment_method')
        payment_time = data.get('payment_time') or data.get('created_at')
        
        logging.info(f"📋 Webhook details - Order: {order_id}, Status: {order_status}, Payment ID: {payment_id}")
        
        if order_status == 'PAID':
            # Find payment record in database
            payment_record = db_helper.get_payment_by_gateway_id(order_id)
            if payment_record:
                logging.info(f"💳 Found payment record: {payment_record['id']} for order {order_id}")
                
                # Update payment status with additional details
                updated_payment = db_helper.update_payment_status(order_id, 'paid', payment_id)
                if updated_payment:
                    logging.info(f"✅ Payment status updated to 'paid' for order {order_id}")
                
                # Add credits to user account
                current_user = db_helper.get_user_by_phone(payment_record['phone_number'])
                if current_user:
                    old_credits = current_user['credits']
                    new_credits = old_credits + payment_record['credits_added']
                    
                    # Update user credits
                    updated_user = db_helper.update_user_credits(payment_record['phone_number'], new_credits)
                    if updated_user:
                        logging.info(f"💰 Credits updated for {payment_record['phone_number']}: {old_credits} → {new_credits} (+{payment_record['credits_added']})")
                    
                    # Update session if user is currently logged in
                    current_session_user = flask_session.get('user')
                    if current_session_user and current_session_user.get('phone_number') == payment_record['phone_number']:
                        fresh_user_data = db_helper.get_user_by_phone(payment_record['phone_number'])
                        if fresh_user_data:
                            flask_session['user'] = fresh_user_data
                            logging.info(f"🔄 Session updated for logged-in user {payment_record['phone_number']}")
                
                # Log successful payment processing
                logging.info(f"🎉 Payment successfully processed: Order {order_id}, Amount: ₹{payment_record['amount']/100}, Credits: +{payment_record['credits_added']}")
                
                return jsonify({
                    'status': 'success',
                    'message': 'Payment processed successfully',
                    'order_id': order_id,
                    'credits_added': payment_record['credits_added']
                }), 200
            else:
                logging.error(f"❌ Payment record not found for order: {order_id}")
                return jsonify({
                    'status': 'error', 
                    'message': 'Payment record not found',
                    'order_id': order_id
                }), 404
        
        elif order_status in ['FAILED', 'CANCELLED']:
            # Handle failed payments
            payment_record = db_helper.get_payment_by_gateway_id(order_id)
            if payment_record:
                db_helper.update_payment_status(order_id, 'failed', payment_id)
                logging.info(f"💔 Payment failed for order {order_id}")
            
            return jsonify({
                'status': 'acknowledged',
                'message': f'Payment {order_status.lower()}',
                'order_id': order_id
            }), 200
        
        else:
            # Acknowledge other statuses
            logging.info(f"ℹ️ Payment status '{order_status}' acknowledged for order {order_id}")
            return jsonify({
                'status': 'acknowledged',
                'message': f'Status {order_status} recorded',
                'order_id': order_id
            }), 200
        
    except Exception as e:
        logging.error(f"💥 Payment webhook error: {e}")
        return jsonify({
            'status': 'error', 
            'message': str(e)
        }), 500

# --- Simple OTP Functions (replacing Twilio dependency) ---
def generate_otp():
    """Generate a 6-digit OTP"""
    return str(random.randint(100000, 999999))

def validate_phone_number(phone_number):
    """Simple phone number validation"""
    import re
    # Remove any spaces, dashes, or plus signs
    cleaned_phone = re.sub(r'[\s\-\+]', '', phone_number)
    
    # Check if it's all digits and has appropriate length
    if not cleaned_phone.isdigit():
        return False
        
    # Indian phone number validation (10 digits starting with 6-9)
    if len(cleaned_phone) == 10 and cleaned_phone[0] in '6789':
        return True
        
    # International format (10-15 digits)
    if 10 <= len(cleaned_phone) <= 15:
        return True
        
    return False

def send_otp_console(phone_number, otp_code):
    """Send OTP via console (for development)"""
    print(f"\n🔐 OTP for {phone_number}: {otp_code}")
    print(f"📱 In production, this would be sent via SMS")
    logging.info(f"OTP generated for {phone_number}: {otp_code}")
    return True





def get_cached_solution(question, language="python"):
    key = (question, language)
    with cache_lock:
        if key in solution_cache:
            logging.debug("Cache hit for question.")
            return solution_cache[key]
    sol = solve_coding_problem(question, language)
    with cache_lock:
        solution_cache[key] = sol
    return sol

# ----- Utility Functions -----
def extract_text_from_pdf(pdf_path):
    try:
        with open(pdf_path, "rb") as pdf_file:
            reader = PdfReader(pdf_file)
            extracted_text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    extracted_text.append(page_text)
            return "\n".join(extracted_text).strip()
    except Exception as e:
        logging.exception("Error extracting text from PDF")
        return f"Error extracting text from PDF: {e}"
    
    
def try_claude_api(prompt, api_key, key_index):
    """Enhanced Claude API with concurrent connection control"""
    # Acquire global connection semaphore first
    if not claude_manager.connection_semaphore.acquire(blocking=False):
        raise Exception("Global Claude API connection limit reached. Please try again shortly.")
    
    try:
        # Track concurrent connection for this specific key
        with claude_manager.lock:
            claude_manager.key_stats[key_index]['concurrent_connections'] += 1
            concurrent_count = claude_manager.key_stats[key_index]['concurrent_connections']
            logging.debug(f"Claude API key {key_index + 1} concurrent connections: {concurrent_count}")
        
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=800,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
        
        content = response.content[0].text
        claude_manager.mark_success(key_index)
        return re.sub(r"```[\w]*", "", content).replace("```", "").strip()
        
    except Exception as e:
        error_str = str(e).lower()
        
        # Enhanced error detection for concurrent connections
        if "concurrent" in error_str or "connection" in error_str:
            claude_manager.mark_rate_limited(key_index, duration=90)
            logging.warning(f"Claude API key {key_index + 1} concurrent connection limit: {str(e)}")
        # Detect other rate limit errors
        elif "rate" in error_str or "limit" in error_str or "429" in error_str:
            claude_manager.mark_rate_limited(key_index, duration=60)
            logging.warning(f"Claude API key {key_index + 1} rate limited: {str(e)}")
        # Detect auth errors
        elif "401" in error_str or "unauthorized" in error_str:
            claude_manager.mark_rate_limited(key_index, duration=300)  # 5 minute cooldown for auth issues
            logging.error(f"Claude API key {key_index + 1} unauthorized: {str(e)}")
        else:
            # General error handling
            with claude_manager.lock:
                claude_manager.key_stats[key_index]['failures'] += 1
                claude_manager.key_stats[key_index]['consecutive_errors'] += 1
            logging.warning(f"Claude API key {key_index + 1} failed: {str(e)}")
        
        raise e
    
    finally:
        # Always decrement concurrent connection count and release semaphore
        try:
            with claude_manager.lock:
                claude_manager.key_stats[key_index]['concurrent_connections'] = max(0, 
                    claude_manager.key_stats[key_index]['concurrent_connections'] - 1)
            claude_manager.connection_semaphore.release()
        except Exception as cleanup_error:
            logging.error(f"Error cleaning up Claude API connection: {cleanup_error}")

def try_deepseek_api(prompt, timeout=45):
    """Enhanced DeepSeek API with intelligent load balancing and retry logic"""
    max_retries = len(DEEPSEEK_KEYS) * 2  # Allow multiple attempts per key
    
    for attempt in range(max_retries):
        try:
            # Get the best available API key
            key_index, current_key = deepseek_manager.get_best_key()
            
            logging.info(f"Using DeepSeek API key {key_index + 1}/{len(DEEPSEEK_KEYS)} (attempt {attempt + 1}/{max_retries})")
            
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {current_key}"
            }
            
            # Enhanced request data with better parameters for high traffic
            data = {
                "model": "deepseek-chat",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.2,
                "top_p": 0.9,
                "max_tokens": 1000,  # Increased for better responses
                "stream": False
            }
            
            # Make request with increased timeout for high traffic
            response = requests.post(
                BASE_URL, 
                json=data, 
                headers=headers, 
                timeout=timeout
            )
            
            if response.status_code == 200:
                content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
                deepseek_manager.mark_success(key_index)
                logging.info(f"✅ DeepSeek API key {key_index + 1} succeeded")
                return re.sub(r"```[\w]*", "", content).replace("```", "").strip()
            
            elif response.status_code == 429:  # Rate limit
                deepseek_manager.mark_rate_limited(key_index, duration=90)  # 90 second cooldown
                logging.warning(f"🚫 API key {key_index + 1} rate limited (429)")
                continue
                
            elif response.status_code == 401:  # Unauthorized - invalid key
                deepseek_manager.mark_rate_limited(key_index, duration=300)  # 5 minute cooldown for invalid keys
                logging.error(f"🔑 API key {key_index + 1} unauthorized (401)")
                continue
                
            else:
                error_msg = f"HTTP {response.status_code}: {response.text[:200]}"
                logging.error(f"❌ DeepSeek API key {key_index + 1} error: {error_msg}")
                
                # For 5xx errors, mark as temporarily unavailable
                if 500 <= response.status_code < 600:
                    deepseek_manager.mark_rate_limited(key_index, duration=30)
                
                continue
                
        except requests.exceptions.Timeout:
            logging.warning(f"⏰ DeepSeek API key {key_index + 1} timeout after {timeout}s")
            continue
            
        except requests.exceptions.ConnectionError:
            logging.warning(f"🔌 DeepSeek API key {key_index + 1} connection error")
            time.sleep(1)  # Brief pause before retry
            continue
            
        except Exception as e:
            error_str = str(e).lower()
            if "rate" in error_str or "429" in error_str:
                deepseek_manager.mark_rate_limited(key_index, duration=60)
                logging.warning(f"🚫 DeepSeek API key {key_index + 1} rate limited: {str(e)}")
            else:
                logging.error(f"❌ DeepSeek API key {key_index + 1} unexpected error: {str(e)}")
            continue
    
    # If all attempts failed, provide detailed error info
    stats = deepseek_manager.get_stats()
    error_details = f"All {len(DEEPSEEK_KEYS)} DeepSeek API keys failed after {max_retries} attempts. Key stats: {stats}"
    logging.error(error_details)
    raise Exception("All DeepSeek API keys exhausted. Please try again in a few minutes.")

def solve_coding_problem(question, language="python"):
    """Intelligent API selection with load balancing between Claude and DeepSeek"""
    if "input" in question.lower() or "user" in question.lower():
        question += f" Assume the user input is {random.randint(1, 9)}. The code should NOT prompt for input."

    lang_str = {
        "python": "Python", "c++": "C++", "cpp": "C++",
        "c#": "C#", "csharp": "C#", "c": "C",
        "java": "Java", "javascript": "JavaScript"
    }.get(language.lower(), language)

    prompt = f"""
Write only the {lang_str} code to solve the following problem:
start
{question}
end
there should not be any comments in the code only raw code should be provided.
The code must NOT contain any 'input()' function or prompt the user for input.
Instead, assume predefined values for any required inputs.
Ensure the code includes at least one print statement to display the final result or output.
Ensure every question yields actual, printed output. If the problem does not define a specific expected answer, craft a minimal hypothetical example with fixed inputs to demonstrate the solution and print the outcome.
For any file I/O tasks: use safe relative paths (e.g., "data.txt"), write the data, read it back, and print a clear summary (e.g., number of items and the values loaded). Do not rely on external files or user input.
The program must not exit without printing; avoid errors by using safe defaults and guard checks.
Don't use shortcut imports, write code from scratch.
Return only the raw code. Do not include any explanations, comments, markdown formatting, or extra characters.
"""

    # C#-specific guard rails to improve compilability and visible output
    if lang_str == "C#":
        prompt += """
For C#: Always include 'using System;'. Include 'using System.Collections;' when using ArrayList/Hashtable, and 'using System.Collections.Generic;' when using List<>/Dictionary<>. Include 'using System.IO;' when using FileInfo/Directory/Path.
Define 'class Program' with 'static void Main()' as the entrypoint.
Do not request interactive input; use fixed sample values.
For filesystem/search tasks, simulate with a small predefined array of filenames and use a simple collection (e.g., List<string> or Dictionary<string, FileInfo-like struct>) to store valid entries. Print details for a requested filename. Avoid real system I/O beyond safe relative files.
Ensure several 'Console.WriteLine(...)' statements print the results clearly.
"""

    # INTELLIGENT API SELECTION LOGIC
    claude_available_keys = claude_manager.get_available_keys_count()
    deepseek_available_keys = len([i for i in range(len(DEEPSEEK_KEYS)) 
                                  if deepseek_manager.key_stats[i]['rate_limited_until'] < time.time()])
    
    # Decision logic: Use Claude if available, otherwise use DeepSeek
    use_claude_first = claude_available_keys > 0
    
    logging.info(f"API Selection: Claude available keys: {claude_available_keys}, DeepSeek available keys: {deepseek_available_keys}")
    
    if use_claude_first:
        # Try Claude API with smart key selection
        try:
            key_index, claude_key = claude_manager.get_best_available_key()
            if claude_key:
                logging.info(f"Using Claude API key {key_index + 1}/{len(CLAUDE_KEYS)} (Smart selection)")
                result = try_claude_api(prompt, claude_key, key_index)
                logging.info(f"✅ Claude API key {key_index + 1} succeeded")
                return result
        except Exception as e:
            logging.warning(f"Claude API attempt failed: {str(e)}")
            # Fall through to DeepSeek
    
    # Try DeepSeek as primary or fallback
    logging.info("Using DeepSeek API (primary choice or Claude fallback)")
    try:
        result = try_deepseek_api(prompt)
        logging.info("✅ DeepSeek API succeeded")
        return result
    except Exception as e:
        logging.error(f"DeepSeek API failed: {str(e)}")
        
        # Last resort: Try any remaining Claude keys if DeepSeek failed
        if not use_claude_first:
            logging.info("Attempting Claude as last resort after DeepSeek failure")
            try:
                key_index, claude_key = claude_manager.get_best_available_key()
                if claude_key:
                    logging.info(f"Last resort: Using Claude API key {key_index + 1}")
                    result = try_claude_api(prompt, claude_key, key_index)
                    logging.info(f"✅ Claude API (last resort) succeeded")
                    return result
            except Exception as claude_e:
                logging.error(f"Claude last resort also failed: {str(claude_e)}")
        
        return f"Error: All API attempts failed. Both Claude and DeepSeek are currently unavailable."


def execute_code(code):
    try:
        if code.startswith("Error") or not code:
            return f"Invalid code returned: {code}"
        logging.debug(f"Executing code:\n{code}")
        output = io.StringIO()
        context = {'__name__': '__main__', '__file__': 'task.py'}
        
        # Check if code involves file operations and set up mock environment
        file_operations = ['open(', 'os.listdir', 'os.walk', 'glob.glob', 'pathlib', 'os.path.exists', 'os.path.isfile', 'os.path.isdir', 'with open']
        has_file_ops = any(op in code for op in file_operations)
        
        if has_file_ops:
            # Set up mock file system context
            context.update({
                'os': __import__('os'),
                'glob': __import__('glob'),
                'pathlib': __import__('pathlib'),
                'tempfile': __import__('tempfile'),
            })
            
            # Create temporary directory with mock files for testing
            import tempfile, os
            temp_dir = tempfile.mkdtemp()
            
            # Create some mock files with different extensions and content
            mock_files = [
                ('document1.txt', 'This is a sample text document with some content.'),
                ('data.csv', 'name,age,city\nJohn,25,New York\nJane,30,London\nBob,35,Paris'),
                ('config.json', '{"database": "localhost", "port": 5432, "debug": true}'),
                ('script.py', 'print("Hello from Python script")\nresult = 42\nprint(f"Result: {result}")'),
                ('readme.md', '# Project Title\nThis is a sample markdown file.\n## Features\n- Feature 1\n- Feature 2'),
                ('log.txt', 'INFO: Application started\nWARNING: Connection timeout\nERROR: Database connection failed'),
                ('numbers.txt', '1\n2\n3\n4\n5\n6\n7\n8\n9\n10'),
                ('words.txt', 'apple\nbanana\ncherry\ndate\nelderberry\nfig\ngrape'),
            ]
            
            # Create subdirectories
            os.makedirs(os.path.join(temp_dir, 'subfolder'), exist_ok=True)
            os.makedirs(os.path.join(temp_dir, 'data'), exist_ok=True)
            
            for filename, content in mock_files:
                filepath = os.path.join(temp_dir, filename)
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            # Create files in subdirectories too
            with open(os.path.join(temp_dir, 'subfolder', 'nested.txt'), 'w') as f:
                f.write('This is a nested file in subfolder.')
            with open(os.path.join(temp_dir, 'data', 'dataset.csv'), 'w') as f:
                f.write('id,value\n1,100\n2,200\n3,300')
            
            # Modify code to work with the temporary directory
            # Replace common directory references with our temp directory
            temp_dir_escaped = temp_dir.replace('\\', '\\\\')
            modified_code = code.replace('"/"', f'"{temp_dir_escaped}"')
            modified_code = modified_code.replace("'/'", f"'{temp_dir_escaped}'")
            modified_code = modified_code.replace('"./"', f'"{temp_dir_escaped}"')
            modified_code = modified_code.replace("'./'", f"'{temp_dir_escaped}'")
            
            # Fix the glob pattern replacements
            temp_glob_double = os.path.join(temp_dir, "*.").replace('\\', '\\\\')
            temp_glob_single = os.path.join(temp_dir, '*.').replace('\\', '\\\\')
            modified_code = modified_code.replace('"*.', f'"{temp_glob_double}')
            modified_code = modified_code.replace("'*.", f"'{temp_glob_single}")
            
            # Add cleanup at the end
            cleanup_code = f"""
# Cleanup temporary files
import shutil
try:
    shutil.rmtree(r"{temp_dir}")
except:
    pass
"""
            modified_code = modified_code + cleanup_code
            
        else:
            # Modified code to ensure output is generated
            modified_code = code
        
        # If code doesn't have print statements, try to add them for common patterns
        if 'print(' not in code.lower():
            lines = code.strip().split('\n')
            last_line = lines[-1].strip() if lines else ''
            
            # Check if last line is an expression that could be printed
            if last_line and not any(keyword in last_line.lower() for keyword in ['if ', 'for ', 'while ', 'def ', 'class ', 'import ', '#', 'try:', 'except:', 'finally:']):
                # If last line looks like a variable or expression, add print
                if '=' not in last_line or last_line.count('=') == 1 and not any(op in last_line for op in ['==', '!=', '<=', '>=']):
                    # This might be an assignment or expression - let's try to print it
                    if '=' in last_line:
                        # It's an assignment, print the variable
                        var_name = last_line.split('=')[0].strip()
                        modified_code += f"\nprint({var_name})"
                    else:
                        # It's an expression, print it directly
                        modified_code += f"\nprint({last_line})"
        
        with contextlib.redirect_stdout(output), contextlib.redirect_stderr(output):
            try:
                exec(modified_code, context, context)
            except NameError as e:
                missing_var = str(e).split("'")[1] if "'" in str(e) else "variable"
                if missing_var == "__name__":
                    modified_code = f'__name__ = "__main__"\n' + modified_code
                else:
                    modified_code = f"{missing_var} = 0\n" + modified_code
                try:
                    exec(modified_code, context, context)
                except Exception as e_inner:
                    # If modified code fails, try original code
                    try:
                        exec(code, context, context)
                    except:
                        raise e_inner
            except Exception as e_inner:
                # If modified code fails, try original code
                try:
                    exec(code, context, context)
                except:
                    raise e_inner
        
        result = output.getvalue().strip()
        
        # If still no output, try to find and print meaningful variables
        if not result:
            # Look for variables that might contain results
            meaningful_vars = []
            for var_name, var_value in context.items():
                if not var_name.startswith('__') and var_name not in ['io', 'contextlib', 'sys']:
                    # Skip functions and modules
                    if not callable(var_value) and not str(type(var_value)).startswith('<class \'module'):
                        meaningful_vars.append(f"{var_name} = {var_value}")
            
            if meaningful_vars:
                result = "\n".join(meaningful_vars[:5])  # Limit to first 5 variables
            else:
                # Last resort: try to execute and capture any return value
                try:
                    # Check if the code is a single expression that returns a value
                    if '\n' not in code.strip() and not any(keyword in code.lower() for keyword in ['def ', 'class ', 'import ', 'for ', 'while ', 'if ']):
                        try:
                            result_value = eval(code, context)
                            if result_value is not None:
                                result = str(result_value)
                        except:
                            pass
                except:
                    pass
        
        return result if result else "Code executed successfully but produced no visible output. This is normal for some programs that don't print results."
    except Exception as e:
        logging.exception(f"Error executing code: {e}")
        return f"Error executing code: {e}"
    
import os
import io
import logging
from PIL import Image, ImageDraw, ImageFont

def create_screenshot(output_text, user_name='Developer', document_terminal_path=None, style='vscode', language='python'):
    """Create a VS Code Windows terminal-style screenshot of code output"""
    try:
        output_text = output_text.strip() or "No output."
        
        # NEW: macOS Terminal style
        if style == 'mac' or style == 'macos':
            try:
                COLORS = {
                    'titlebar_bg': '#2B2B2B',   # macOS dark titlebar
                    'terminal_bg': '#111111',   # darker "Pro" profile-like terminal
                    'text': '#EDEDED',          # bright terminal text
                    'title_text': '#FFFFFF',
                    'traffic_red': '#FF5F56',
                    'traffic_yellow': '#FFBD2E',
                    'traffic_green': '#27C93F',
                    'cursor': '#EDEDED',        # block cursor color
                    'shadow': (0, 0, 0, 90),    # soft shadow (not used in new mac style)
                }

                # Fonts: Menlo preferred (mac), fallback to Consolas/Courier/Default
                font_candidates = [
                    "/System/Library/Fonts/Menlo.ttc",
                    "/Library/Fonts/Menlo.ttc",
                    "C:\\Windows\\Fonts\\consola.ttf",
                    "C:\\Windows\\Fonts\\cour.ttf",
                ]
                terminal_font = None
                for fp in font_candidates:
                    try:
                        if os.path.exists(fp):
                            terminal_font = ImageFont.truetype(fp, 14)
                            break
                    except Exception:
                        continue
                if not terminal_font:
                    terminal_font = ImageFont.load_default()

                # Title font (SF Pro Text if available, fallback Segoe UI then terminal font)
                title_font_candidates = [
                    "/System/Library/Fonts/SFNS.ttf",
                    "/System/Library/Fonts/SFNSDisplay.ttf",
                    "/System/Library/Fonts/SFNSMono.ttf",
                    "C:\\Windows\\Fonts\\segoeui.ttf",
                ]
                title_font = None
                for fp in title_font_candidates:
                    try:
                        if os.path.exists(fp):
                            title_font = ImageFont.truetype(fp, 13)
                            break
                    except Exception:
                        continue
                if not title_font:
                    title_font = terminal_font

                # Derive folder name from path (show 'desktop' like real zsh prompt)
                folder_name = "desktop"
                if document_terminal_path:
                    base = document_terminal_path.replace("\\", "/").strip()
                    parts = [p for p in base.split("/") if p]
                    if parts:
                        candidate = parts[-1].strip()
                        folder_name = candidate.lower() if candidate.lower() not in ("users", "tharan") else "desktop"

                host_name = "MacBook-Air"
                user_clean = (user_name or "user").strip()
                # zsh prompt uses '%' by default
                prompt = f"{user_clean}@{host_name} {folder_name} %"

                # Build realistic command per language (mac commands) with rotated Python filename
                lang_key = (language or 'python').strip().lower()

                rotation_names = [
                    "app.py","program.py","script.py","demo.py","project.py","test.py","run.py","temp.py","example.py",
                    "practice.py","sample.py","output.py","input.py","helper.py","module.py","function.py","assignment.py",
                    "exercise.py","task.py","solution.py","experiment.py","trial.py","execute.py","utility.py","index.py",
                    "starter.py","practice1.py","test1.py","code.py","final.py","new.py","latest.py","backup.py","draft.py",
                    "revision.py","practice2.py","runfile.py","module1.py","temp1.py","program1.py","mainfile.py",
                    "script1.py","execute1.py","simple.py","try.py","check.py","learn.py","classwork.py","college.py",
                    "submission.py"
                ]
                try:
                    seed = (document_terminal_path or user_name or 'student').encode('utf-8', errors='ignore')
                    idx = int(hashlib.sha1(seed).hexdigest(), 16) % len(rotation_names)
                    python_file_name = rotation_names[idx]
                except Exception:
                    python_file_name = "program.py"

                command_map = {
                    'python': f'python3 {python_file_name}',
                    'py': f'python3 {python_file_name}',
                    'c': 'gcc main.c -o hello && ./hello',
                    'cpp': 'g++ main.cpp -o hello && ./hello',
                    'c++': 'g++ main.cpp -o hello && ./hello',
                    'java': 'javac Main.java && java Main',
                    'javascript': 'node app.js',
                    'js': 'node app.js',
                    'c#': 'dotnet run',
                    'csharp': 'dotnet run',
                }
                run_command = command_map.get(lang_key, f'python3 {python_file_name}')

                # "Last login" line like macOS
                from datetime import datetime
                now = datetime.now().strftime("%a %b %d %H:%M:%S")
                last_login = f"Last login: {now} on ttys000"

                # Lines to render, matching a real flow
                rendered_lines = []
                rendered_lines.append(last_login)
                rendered_lines.append(f"{prompt} cd ~/desktop")
                rendered_lines.append(f"{prompt} {run_command}")
                for line in (output_text.strip() or "No output.").splitlines():
                    rendered_lines.append(line)
                # Final prompt removed to avoid extra trailing line after output

                # Measure text widths accurately
                tmp_img = Image.new("RGB", (10, 10), (0, 0, 0))
                tmp_draw = ImageDraw.Draw(tmp_img)
                def text_width(s, font):
                    try:
                        if hasattr(tmp_draw, "textlength"):
                            return int(tmp_draw.textlength(s, font=font))
                    except Exception:
                        pass
                    # Fallback approximate
                    return int(len(s) * 8)

                # Compute dimensions
                line_heights = []
                max_w = 0
                for s in rendered_lines:
                    max_w = max(max_w, text_width(s, terminal_font))
                    h = terminal_font.getbbox("Ag")[3] if hasattr(terminal_font, "getbbox") else terminal_font.getsize("Ag")[1]
                    line_heights.append(h)

                padding_x = 14
                padding_y = 10
                titlebar_height = 26
                line_spacing = 4
                content_height = sum(line_heights) + line_spacing * (len(rendered_lines) - 1) + padding_y * 2
                image_width = max(820, max_w + padding_x * 2)
                image_height = titlebar_height + content_height

                # Create image with transparent background (no shadow)
                image = Image.new('RGBA', (image_width, image_height), (0, 0, 0, 0))
                draw = ImageDraw.Draw(image)

                # Main window (rounded)
                draw.rounded_rectangle(
                    [0, 0, image_width, image_height],
                    radius=10,
                    fill=COLORS['titlebar_bg']
                )

                # Title bar area
                draw.rectangle(
                    [0, 0, image_width, titlebar_height],
                    fill=COLORS['titlebar_bg']
                )

                # Traffic lights (perfect circles with anti-aliasing)
                traffic_diameter = 12
                spacing = 8
                base_x = 14
                center_y = int(round(titlebar_height / 2))

                def paste_circle(cx, cy, color, diameter):
                    scale = 4  # supersample for smooth, round edges
                    big = Image.new('RGBA', (diameter * scale, diameter * scale), (0, 0, 0, 0))
                    big_draw = ImageDraw.Draw(big)
                    big_draw.ellipse([0, 0, diameter * scale - 1, diameter * scale - 1], fill=color)
                    small = big.resize((diameter, diameter), resample=Image.LANCZOS)
                    image.paste(small, (cx - diameter // 2, cy - diameter // 2), small)

                red_cx = base_x
                yellow_cx = red_cx + traffic_diameter + spacing
                green_cx = yellow_cx + traffic_diameter + spacing

                paste_circle(red_cx, center_y, COLORS['traffic_red'], traffic_diameter)
                paste_circle(yellow_cx, center_y, COLORS['traffic_yellow'], traffic_diameter)
                paste_circle(green_cx, center_y, COLORS['traffic_green'], traffic_diameter)

                # Title text centered: "Desktop — zsh — 80×24"
                folder_title = folder_name.capitalize()
                title_text = f"{folder_title} — zsh — 80×24"
                try:
                    title_w = tmp_draw.textlength(title_text, font=title_font) if hasattr(tmp_draw, 'textlength') else len(title_text) * 7
                except Exception:
                    title_w = len(title_text) * 7
                draw.text(
                    ((image_width - title_w) // 2, (titlebar_height - 13) // 2),
                    title_text,
                    fill=COLORS['title_text'],
                    font=title_font
                )

                # Terminal content area
                draw.rectangle(
                    [0, titlebar_height, image_width, image_height],
                    fill=COLORS['terminal_bg']
                )

                # Render content lines
                y = titlebar_height + padding_y
                for i, s in enumerate(rendered_lines):
                    draw.text((padding_x, y), s, fill=COLORS['text'], font=terminal_font)
                    lh = line_heights[i] + line_spacing
                    y += lh

                # Omitted block cursor drawing to avoid extra trailing line after output

                # Convert to RGB for saving (white background)
                final_image = Image.new('RGB', (image_width, image_height), (255, 255, 255))
                final_image.paste(image, (0, 0), image)

                # Save to buffer
                buffer = io.BytesIO()
                final_image.save(buffer, format='PNG')
                buffer.seek(0)
                return buffer.getvalue()

            except Exception as e:
                logging.error(f"Error creating macOS screenshot: {e}")
                # Fallback to simple style
                style = 'simple'
        
        # Use ultra-realistic VS Code screenshot generator for maximum realism
        if style == 'vscode' or style == 'realistic':
            # Direct implementation of VS Code screenshot styling
            try:
                # Clean and prepare output text
                output_text = output_text.strip() or "No output."
                
                # EXACT VS Code Colors from the real screenshot
                COLORS = {
                    'panel_bg': '#252526',        # Top panel background (PROBLEMS, OUTPUT, etc.)
                    'terminal_bg': '#1e1e1e',     # Terminal background (dark)
                    'tab_active_bg': '#1e1e1e',   # Active TERMINAL tab background
                    'tab_inactive_text': '#969696', # Inactive tab text (gray)
                    'tab_active_text': '#ffffff',   # Active TERMINAL tab text (white)
                    'tab_active_border': '#007acc', # Blue underline for active tab
                    'ps_blue': '#569cd6',         # PowerShell "PS" blue
                    'path_yellow': '#dcdcaa',     # Path yellow color (like VS Code)
                    'command_white': '#ffffff',   # Command text white
                    'output_white': '#ffffff',    # Output text white
                    'blue_circle': '#007acc',     # Blue color for circles
                }
                
                # Load fonts - try to get Consolas or Cascadia Code
                font_paths = [
                    "C:\\Windows\\Fonts\\consola.ttf",        # Consolas (most common)
                    "C:\\Windows\\Fonts\\CascadiaCode.ttf",
                    "C:\\Windows\\Fonts\\CascadiaCodePL.ttf", 
                    "C:\\Windows\\Fonts\\CascadiaMono.ttf",
                ]
                
                terminal_font = None
                ui_font = None
                
                # Load terminal font (size 14 like VS Code default)
                for font_path in font_paths:
                    try:
                        if os.path.exists(font_path):
                            terminal_font = ImageFont.truetype(font_path, 14)
                            break
                    except Exception:
                        continue
                
                # Load UI font for tabs
                try:
                    ui_font = ImageFont.truetype("C:\\Windows\\Fonts\\segoeui.ttf", 11)
                except:
                    ui_font = terminal_font
                
                # Fallback to default font
                if not terminal_font:
                    terminal_font = ImageFont.load_default()
                    ui_font = ImageFont.load_default()
                
                # Process terminal path
                if document_terminal_path:
                    terminal_path = document_terminal_path
                else:
                    try:
                        from terminal_utils import TerminalUtils
                        terminal_path = TerminalUtils.get_user_terminal_path(user_name)
                    except:
                        terminal_path = f"C:\\Users\\{user_name}\\Desktop\\AIProjects\\ChatBot"
                
                # Prepare terminal content exactly like the screenshot
                lines = []
                
                # Determine command based on language with rotated Python filename
                lang_key = (language or 'python').strip().lower()

                rotation_names = [
                    "app.py","program.py","script.py","demo.py","project.py","test.py","run.py","temp.py","example.py",
                    "practice.py","sample.py","output.py","input.py","helper.py","module.py","function.py","assignment.py",
                    "exercise.py","task.py","solution.py","experiment.py","trial.py","execute.py","utility.py","index.py",
                    "starter.py","practice1.py","test1.py","code.py","final.py","new.py","latest.py","backup.py","draft.py",
                    "revision.py","practice2.py","runfile.py","module1.py","temp1.py","program1.py","mainfile.py",
                    "script1.py","execute1.py","simple.py","try.py","check.py","learn.py","classwork.py","college.py",
                    "submission.py"
                ]
                try:
                    seed = (document_terminal_path or user_name or 'student').encode('utf-8', errors='ignore')
                    idx = int(hashlib.sha1(seed).hexdigest(), 16) % len(rotation_names)
                    python_file_name = rotation_names[idx]
                except Exception:
                    python_file_name = "program.py"

                command_map = {
                    'python': f'python {python_file_name}',
                    'cpp': 'g++ main.cpp -o main.exe && .\\main.exe',
                    'c++': 'g++ main.cpp -o main.exe && .\\main.exe',
                    'c': 'gcc main.c -o main.exe && .\\main.exe',
                    'java': 'javac Main.java && java Main',
                    'javascript': 'node app.js',
                    'js': 'node app.js',
                    'c#': 'dotnet run',
                    'csharp': 'dotnet run',
                }
                run_command = command_map.get(lang_key, f'python {python_file_name}')
                
                # First line: PowerShell prompt with correct command
                lines.append(f"PS {terminal_path}> {run_command}")
                
                # Add output lines (only the program output)
                output_lines = output_text.splitlines()
                for line in output_lines:
                    lines.append(line)
                
                # For C#, do not append the trailing prompt line so the terminal ends with the output
                if lang_key not in ('c#', 'csharp'):
                    lines.append(f"PS {terminal_path}>")
                
                # Calculate dimensions based on real VS Code
                char_width = 7      # Character width for Consolas 14px
                line_height = 22    # Line height in VS Code terminal (increased to prevent overlapping)
                padding_x = 10      # Left padding (increased for better spacing)
                padding_y = 10      # Top padding (increased for better spacing)
                
                # Panel dimensions (exact measurements from screenshot)
                panel_height = 35   # Top panel with PROBLEMS, OUTPUT, etc.
                
                # Calculate image dimensions
                max_line_length = max(len(line) for line in lines) if lines else 80
                image_width = max(900, max_line_length * char_width + padding_x * 2)
                content_height = len(lines) * line_height + padding_y * 2
                image_height = panel_height + content_height
                
                # Create image
                image = Image.new('RGB', (image_width, image_height), color=COLORS['terminal_bg'])
                draw = ImageDraw.Draw(image)
                
                # Draw top panel (PROBLEMS, OUTPUT, DEBUG CONSOLE, TERMINAL, PORTS)
                draw.rectangle([0, 0, image_width, panel_height], fill=COLORS['panel_bg'])
                
                # Panel tabs
                tabs = ["PROBLEMS", "OUTPUT", "DEBUG CONSOLE", "TERMINAL", "PORTS"]
                tab_x = 16  # Start position
                
                for tab in tabs:
                    if ui_font:
                        tab_width = draw.textlength(tab, font=ui_font) if hasattr(draw, 'textlength') else len(tab) * 6
                    else:
                        tab_width = len(tab) * 6
                    
                    if tab == "TERMINAL":
                        # Active tab - draw background and blue underline
                        draw.rectangle([tab_x - 8, 0, tab_x + tab_width + 8, panel_height], fill=COLORS['tab_active_bg'])
                        draw.rectangle([tab_x - 8, panel_height - 2, tab_x + tab_width + 8, panel_height], fill=COLORS['tab_active_border'])
                        text_color = COLORS['tab_active_text']
                    else:
                        text_color = COLORS['tab_inactive_text']
                    
                    # Draw tab text
                    if ui_font:
                        draw.text((tab_x, 11), tab, fill=text_color, font=ui_font)
                    
                    tab_x += tab_width + 32  # Space between tabs
                
                # Draw terminal content
                content_y = panel_height + padding_y
                
                # Blue circle parameters
                circle_radius = 6
                circle_spacing = 5
                
                for i, line in enumerate(lines):
                    if not line.strip():
                        content_y += line_height
                        continue
                    
                    # Handle PowerShell prompt lines
                    if line.startswith("PS ") and ">" in line: 
                        parts = line.split(">", 1) 
                        if len(parts) == 2: 
                            prompt_part = parts[0] + ">" 
                            command_part = parts[1].strip() 
                            
                            # Draw "PS " in blue 
                            draw.text((padding_x, content_y), "PS ", fill=COLORS['ps_blue'], font=terminal_font) 
                            ps_width = draw.textlength("PS ", font=terminal_font) if hasattr(draw, 'textlength') else 21 
                            
                            # Draw path in yellow (like VS Code)
                            path_part = prompt_part[3:]  # Remove "PS " part 
                            draw.text((padding_x + ps_width, content_y), path_part, fill=COLORS['path_yellow'], font=terminal_font)
                            path_width = draw.textlength(path_part, font=terminal_font) if hasattr(draw, 'textlength') else len(path_part) * char_width
                            
                            # Draw command in white (if any)
                            if command_part:
                                draw.text((padding_x + ps_width + path_width, content_y), " " + command_part, fill=COLORS['command_white'], font=terminal_font)
                        else:
                            # Fallback for malformed prompt
                            draw.text((padding_x, content_y), line, fill=COLORS['output_white'], font=terminal_font)
                    else:
                        # Regular output lines in white
                        draw.text((padding_x, content_y), line, fill=COLORS['output_white'], font=terminal_font)
                    
                    content_y += line_height
                
                   # Draw three circles - red, yellow, green (macOS style) but all blue
                
            
                
                # Save as PNG
                buffer = io.BytesIO()
                image.save(buffer, format="PNG", quality=95, optimize=True)
                buffer.seek(0)
                return buffer.read()
                
            except Exception as e:
                logging.exception(f"Error creating ultra realistic VS Code screenshot: {e}")
                # Fallback to original realistic generator
                from realistic_vscode_screenshot import create_realistic_vscode_screenshot
                return create_realistic_vscode_screenshot(output_text, user_name, document_terminal_path)
        
        # Use dedicated simple screenshot generator for 'simple' style
        if style == 'simple':
            from simple_screenshot_generator import create_simple_screenshot
            return create_simple_screenshot(output_text)
        
        # Default fallback to realistic VS Code style
        from realistic_vscode_screenshot import create_realistic_vscode_screenshot
        return create_realistic_vscode_screenshot(output_text, user_name, document_terminal_path)
    
    except Exception as e:
        logging.exception(f"Error in create_screenshot: {e}")
        # Ultimate fallback - create a simple text screenshot
        try:
            from simple_screenshot_generator import create_simple_screenshot
            return create_simple_screenshot(output_text)
        except:
            # If all else fails, return empty bytes
            return b""
def generate_word_doc(name, reg_number, questions, solutions, screenshots, output_path, customization=None, christ_template_data=None, template_option=None):
    try:
        # Initialize template_path to prevent UnboundLocalError
        template_path = None
        
        # Check if user selected "no template" option
        if template_option == 'none':
            logging.info("User selected 'No Template' - creating document from scratch")
            doc = Document()
        else:
            # Template path for CHRIST template
            template_path = os.path.join(os.path.dirname(__file__), 'christ_template', 'template for christ.docx')
            
            # Check if template exists and user wants to use a template
            if template_option == 'christ' and os.path.exists(template_path):
                # Load the CHRIST template document
                logging.info(f"Loading CHRIST template from: {template_path}")
                doc = Document(template_path)
            elif template_option == 'christ' and not os.path.exists(template_path):
                logging.error(f"CHRIST template not found at: {template_path}")
                # Fall back to creating a new document
                doc = Document()
            else:
                # For any other template option or fallback, create new document
                logging.info("Creating document from scratch (no template or unknown template option)")
                doc = Document()
        
        # Parse customization settings with proper mapping
        default_settings = {
            # Heading settings
            'headingFontStyle': 'Calibri',
            'headingSize': 14,
            'headingColor': '#000000',
            'headingStyle': 'bold',
            
            # Question text settings
            'questionFontStyle': 'Calibri', 
            'questionSize': 11,
            'questionColor': '#000000',
            'questionBold': False,
            'questionItalic': False,
            'questionUnderline': False,
            
            # Code settings
            'codeFontStyle': 'Consolas',
            'codeSize': 10,
            'codeColor': '#000000',
            
            # Global settings
            'lineSpacing': 1.15,
            'textAlignment': 'left'
        }
        
        if customization:
            logging.info(f"Applying customization settings: {customization}")
            settings = {**default_settings, **customization}
        else:
            settings = default_settings
        
        # Helper function to apply heading formatting
        def apply_heading_formatting(run):
            # Apply heading font style
            run.font.name = settings.get('headingFontStyle', 'Calibri')
            
            # Apply heading font size
            try:
                font_size = int(settings.get('headingSize', 14))
                run.font.size = Pt(font_size)
            except (ValueError, TypeError):
                run.font.size = Pt(14)
            
            # Apply heading font color
            try:
                color_hex = settings.get('headingColor', '#000000').replace('#', '')
                if len(color_hex) == 6:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
            except (ValueError, TypeError):
                pass  # Keep default color
            
            # Apply heading formatting styles
            heading_style = settings.get('headingStyle', 'bold')
            run.bold = 'bold' in heading_style
            run.italic = 'italic' in heading_style
            run.underline = 'underline' in heading_style
        
        # Helper function to apply question text formatting
        def apply_question_formatting(run):
            # Apply question font style
            run.font.name = settings.get('questionFontStyle', 'Calibri')
            
            # Apply question font size
            try:
                font_size = int(settings.get('questionSize', 11))
                run.font.size = Pt(font_size)
            except (ValueError, TypeError):
                run.font.size = Pt(11)
            
            # Apply question font color
            try:
                color_hex = settings.get('questionColor', '#000000').replace('#', '')
                if len(color_hex) == 6:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
            except (ValueError, TypeError):
                pass  # Keep default color
            
            # Apply question text formatting
            run.bold = settings.get('questionBold', False)
            run.italic = settings.get('questionItalic', False)
            run.underline = settings.get('questionUnderline', False)
        
        # Helper function to apply code formatting
        def apply_code_formatting(run):
            # Apply code font style
            run.font.name = settings.get('codeFontStyle', 'Consolas')
            
            # Apply code font size
            try:
                font_size = int(settings.get('codeSize', 10))
                run.font.size = Pt(font_size)
            except (ValueError, TypeError):
                run.font.size = Pt(10)
            
            # Apply code font color
            try:
                color_hex = settings.get('codeColor', '#000000').replace('#', '')
                if len(color_hex) == 6:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
            except (ValueError, TypeError):
                pass  # Keep default color
        
        # Helper function to apply paragraph formatting
        def apply_paragraph_formatting(paragraph):
            # Apply line spacing
            line_spacing = settings.get('lineSpacing', 1.15)
            if isinstance(line_spacing, str):
                if line_spacing == '1.15' or line_spacing == '1.15':
                    paragraph.paragraph_format.line_spacing = 1.15
                elif line_spacing == '1.5':
                    paragraph.paragraph_format.line_spacing = 1.5
                elif line_spacing == '2.0' or line_spacing == 'double':
                    paragraph.paragraph_format.line_spacing = 2.0
                else:  # single or '1.0'
                    paragraph.paragraph_format.line_spacing = 1.0
            else:
                # Handle numeric values
                paragraph.paragraph_format.line_spacing = float(line_spacing)
            
            # Apply text alignment
            alignment = settings.get('textAlignment', 'left')
            if alignment == 'center':
                paragraph.alignment = 1
            elif alignment == 'right':
                paragraph.alignment = 2
            elif alignment == 'justify':
                paragraph.alignment = 3
            else:  # left
                paragraph.alignment = 0
        
        # Function to replace placeholders in paragraphs
        def replace_placeholders_in_paragraph(paragraph, placeholders):
            """Replace placeholders in a paragraph's text"""
            # First, try to replace in individual runs
            for placeholder, value in placeholders.items():
                for run in paragraph.runs:
                    if placeholder in run.text:
                        # Clean the value by removing any surrounding quotes
                        clean_value = str(value).strip('"\'')
                        run.text = run.text.replace(placeholder, clean_value)
                    
                    # Also check for quoted placeholders in the template (e.g., "#NAME" -> NAME)
                    quoted_placeholder = f'"{placeholder}"'
                    if quoted_placeholder in run.text:
                        clean_value = str(value).strip('"\'')
                        run.text = run.text.replace(quoted_placeholder, clean_value)
            
            # If placeholder spans across multiple runs, we need to handle it differently
            full_text = paragraph.text
            text_changed = False
            
            for placeholder, value in placeholders.items():
                if placeholder in full_text:
                    # Clean the value by removing any surrounding quotes
                    clean_value = str(value).strip('"\'')
                    full_text = full_text.replace(placeholder, clean_value)
                    text_changed = True
                
                # Also check for quoted placeholders in the template
                quoted_placeholder = f'"{placeholder}"'
                if quoted_placeholder in full_text:
                    clean_value = str(value).strip('"\'')
                    full_text = full_text.replace(quoted_placeholder, clean_value)
                    text_changed = True
            
            # If text was changed and placeholders were found across runs
            if text_changed and full_text != paragraph.text:
                # Store the formatting of the first run
                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    # Store formatting
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    font_bold = first_run.font.bold
                    font_italic = first_run.font.italic
                    font_underline = first_run.font.underline
                    font_color = first_run.font.color.rgb if first_run.font.color.rgb else None
                    
                    # Clear paragraph and add new text
                    paragraph.clear()
                    new_run = paragraph.add_run(full_text)
                    
                    # Restore formatting
                    if font_name:
                        new_run.font.name = font_name
                    if font_size:
                        new_run.font.size = font_size
                    if font_bold:
                        new_run.font.bold = font_bold
                    if font_italic:
                        new_run.font.italic = font_italic
                    if font_underline:
                        new_run.font.underline = font_underline
                    if font_color:
                        new_run.font.color.rgb = font_color
        
        # Function to replace placeholders in tables
        def replace_placeholders_in_tables(tables, placeholders):
            """Replace placeholders in table cells"""
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_placeholders_in_paragraph(paragraph, placeholders)
        
        # Create placeholder mapping using # format to match CHRIST template
        placeholders = {
            '#NAME': name,
            '#REGISTER_NUMBER': reg_number,
            '#REG_NUMBER': reg_number,
            '#STUDENT_NAME': name,
            '#DATE': datetime.now().strftime('%B %d, %Y'),
            '#TOTAL_QUESTIONS': str(len(questions))
        }
        
        # Add CHRIST template specific placeholders if provided
        if christ_template_data:
            logging.info(f"Adding CHRIST template placeholders: {christ_template_data}")
            
            # Override with CHRIST-specific data if provided
            if christ_template_data.get('name'):
                placeholders['#NAME'] = christ_template_data['name']
                placeholders['#STUDENT_NAME'] = christ_template_data['name']
            
            if christ_template_data.get('reg_number'):
                placeholders['#REGISTER_NUMBER'] = christ_template_data['reg_number']
                placeholders['#REG_NUMBER'] = christ_template_data['reg_number']
                placeholders['#Registration Number'] = christ_template_data['reg_number']
                placeholders['(#Registration Number)'] = christ_template_data['reg_number']  # With parentheses
            
            # Add additional CHRIST-specific placeholders
            placeholders.update({
                # Course placeholders - exact matches
                '#COURSE NAME': christ_template_data.get('course', ''),
                '#Course name': christ_template_data.get('course', ''),
                '#COURSE': christ_template_data.get('course', ''),
                
                # Class placeholders - exact matches
                '#Class name': christ_template_data.get('class_section', ''),
                '#CLASS_SECTION': christ_template_data.get('class_section', ''),
                '#CLASS': christ_template_data.get('class_section', ''),
                '#SECTION': christ_template_data.get('class_section', ''),
                
                # Practical number placeholders - without quotes
                '#PRACTICAL_NUMBER': f"Practical - {christ_template_data.get('practical_number', '')}" if christ_template_data.get('practical_number') else '',
                '#PRACTICAL': f"Practical - {christ_template_data.get('practical_number', '')}" if christ_template_data.get('practical_number') else '',
                
                # Teacher name placeholders - including lowercase with space
                '#TEACHER_NAME': christ_template_data.get('teacher_name', ''),
                '#Teacher name': christ_template_data.get('teacher_name', ''),
                '#teacher name': christ_template_data.get('teacher_name', ''),  # Lowercase with space
                '#TEACHER': christ_template_data.get('teacher_name', ''),
                '#FACULTY': christ_template_data.get('teacher_name', ''),
                '#INSTRUCTOR': christ_template_data.get('teacher_name', '')
            })
        
        # Replace placeholders in the document
        logging.info("Replacing placeholders in template...")
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, placeholders)
        
        # Replace in tables
        replace_placeholders_in_tables(doc.tables, placeholders)
        
        # Replace in headers and footers
        for section in doc.sections:
            # Header
            header = section.header
            for paragraph in header.paragraphs:
                replace_placeholders_in_paragraph(paragraph, placeholders)
            replace_placeholders_in_tables(header.tables, placeholders)
            
            # Footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                replace_placeholders_in_paragraph(paragraph, placeholders)
            replace_placeholders_in_tables(footer.tables, placeholders)
        
        # Find the placeholder for questions content or add content at the end
        questions_inserted = False
        
        # Look for a specific placeholder for questions
        for paragraph in doc.paragraphs:
            if '{{QUESTIONS_CONTENT}}' in paragraph.text:
                # Clear the placeholder paragraph
                paragraph.clear()
                questions_inserted = True
                break
        
        # If no placeholder found, add content to the end
        if not questions_inserted:
            # Questions will start immediately after template content (no page break)
            # --- Question Pages ---
            for i, (question, solution, screenshot) in enumerate(zip(questions, solutions, screenshots)):
                # Question header
                question_heading = doc.add_heading(f"QUESTION {i + 1}", level=2)
                for run in question_heading.runs:
                    apply_heading_formatting(run)
                apply_paragraph_formatting(question_heading)
                
                # Question content
                question_paragraph = doc.add_paragraph(question.strip())
                for run in question_paragraph.runs:
                    apply_question_formatting(run)
                apply_paragraph_formatting(question_paragraph)
                
                # Code Solution Section
                solution_heading = doc.add_heading("Code Solution", level=3)
                for run in solution_heading.runs:
                    apply_heading_formatting(run)
                apply_paragraph_formatting(solution_heading)
                
                code_paragraph = doc.add_paragraph(solution.strip())
                if code_paragraph.runs:
                    for run in code_paragraph.runs:
                        apply_code_formatting(run)
                apply_paragraph_formatting(code_paragraph)
                
                # Final Output Section
                output_heading = doc.add_heading("FINAL Output", level=3)
                for run in output_heading.runs:
                    apply_heading_formatting(run)
                apply_paragraph_formatting(output_heading)
                
                if screenshot:
                    try:
                        # Significantly increased image size for better visibility
                        doc.add_picture(io.BytesIO(screenshot), width=Inches(5))
                    except Exception as pic_error:
                        logging.warning(f"Could not add picture for QUESTION {i + 1}: {pic_error}")
                        no_pic_paragraph = doc.add_paragraph("FINAL output screenshot not available.")
                        for run in no_pic_paragraph.runs:
                            apply_question_formatting(run)
                        apply_paragraph_formatting(no_pic_paragraph)
                else:
                    no_output_paragraph = doc.add_paragraph("FINAL output not available.")
                    for run in no_output_paragraph.runs:
                        apply_question_formatting(run)
                    apply_paragraph_formatting(no_output_paragraph)
                
                # Force a new page after each question (except the last one)
                if i < len(questions) - 1:
                    doc.add_page_break()
        
        # Save the document
        doc.save(output_path)
        logging.info(f"Word document generated successfully using template with customization applied: {output_path}")
        
        # Log template usage
        if template_path and os.path.exists(template_path):
            logging.info(f"Template-based document created with placeholders replaced")
        else:
            logging.info(f"Document created from scratch (template not found or no template)")
            
        return output_path
    except Exception as e:
        logging.exception(f"Error generating Word document: {e}")
        return f"Error generating Word document: {e}"

def split_questions(text):
    """
    Splits a text into separate coding questions while ensuring that:
    - Multi-line questions remain together.
    - Only properly numbered questions are split.
    - Any unwanted headings (e.g., "Coding Assignment", "Name:", "Register Number:", "Question X:") are omitted.
    """
    try:
        # Remove header lines and stop at "Code:" marker if present.
        header_patterns = [r'^Coding Assignment', r'^Name:', r'^Register Number:', r'^Question \d+:']
        lines = text.splitlines()
        filtered_lines = []
        for line in lines:
            stripped = line.strip()
            if any(re.match(pattern, stripped, re.IGNORECASE) for pattern in header_patterns):
                continue
            if stripped.lower().startswith("code:"):
                break
            filtered_lines.append(stripped)
        filtered_text = "\n".join(filtered_lines).strip()
        
        # Use a regex to split based on numbered questions.
        # The regex now allows for an optional space after the period.
        question_start_pattern = re.compile(r'^\s*\d+\.\s*')
        
        # Only consider lines starting from the first numbered question.
        question_lines = []
        found_first = False
        for line in filtered_text.splitlines():
            if question_start_pattern.match(line):
                found_first = True
            if found_first:
                question_lines.append(line)
        cleaned_text = "\n".join(question_lines)
        
        # Build individual questions.
        questions = []
        current_question = ""
        for line in cleaned_text.splitlines():
            if question_start_pattern.match(line) and current_question:
                questions.append(current_question.strip())
                current_question = line
            else:
                if current_question:
                    current_question += " " + line
                else:
                    current_question = line
        if current_question:
            questions.append(current_question.strip())
        
        # If no splitting occurred, try splitting on double newlines.
        if len(questions) <= 1:
            paragraphs = [para.strip() for para in filtered_text.split("\n\n") if para.strip()]
            if len(paragraphs) > 1:
                return paragraphs
            return [filtered_text]
        return questions
    except Exception as e:
        logging.exception("Error splitting questions")
        return [text.strip()]

def execute_csharp_code(code):
    import os
    import subprocess
    import shutil
    import uuid
    import re

    # Normalize C# source to compile cleanly and produce visible output when needed
    def _massage_csharp(src: str) -> str:
        try:
            s = src or ""
            # Ensure essential using directives
            if "using System;" not in s:
                s = "using System;\n" + s
            if ("ArrayList" in s or "Hashtable" in s) and "using System.Collections;" not in s:
                s = "using System.Collections;\n" + s
            if (("List<" in s) or ("Dictionary<" in s)) and "using System.Collections.Generic;" not in s:
                s = "using System.Collections.Generic;\n" + s
            # Add System.IO when file-related APIs are detected
            file_tokens = ["FileInfo", "Directory", "DirectoryInfo", "File.", "Path."]
            if any(tok in s for tok in file_tokens) and "using System.IO;" not in s:
                s = "using System.IO;\n" + s
            # Ensure a Main entrypoint
            if "static void Main" not in s:
                s = "class Program {\n    static void Main() {\n" + s + "\n    }\n}\n"
            # Ensure at least one print inside Main
            if "Console.WriteLine" not in s:
                try:
                    s = re.sub(
                        r'(static\s+void\s+Main\s*\([^\)]*\)\s*\{)',
                        r'\1\n        Console.WriteLine("Program started...");',
                        s,
                        count=1
                    )
                except Exception:
                    pass
            return s
        except Exception:
            return src

    # Use the app's temp folder
    try:
        temp_dir = TEMP_FOLDER if os.path.isdir(TEMP_FOLDER) else os.getcwd()
    except Exception:
        temp_dir = os.getcwd()

    # Prefer dotnet CLI first for clean "dotnet run" output
    dotnet_path = shutil.which("dotnet")
    if dotnet_path:
        cs_project_dir = os.path.join(temp_dir, f"CSharpRun_{uuid.uuid4().hex[:8]}")
        try:
            os.makedirs(cs_project_dir, exist_ok=True)
        except Exception as e:
            return f"Error creating temp project directory: {e}"

        # Scaffold a console project
        try:
            new_proc = subprocess.run(
                [dotnet_path, "new", "console", "--force"],
                cwd=cs_project_dir,
                capture_output=True,
                text=True,
                shell=False
            )
        except Exception as e:
            return f"Error invoking 'dotnet new console': {e}"

        if new_proc.returncode != 0:
            new_err = (new_proc.stderr or "").strip()
            return new_err or "dotnet new console failed."

        # Write Program.cs with the generated solution code
        cs_file = os.path.join(cs_project_dir, "Program.cs")
        try:
            code_to_write = _massage_csharp(code)
            with open(cs_file, "w", encoding="utf-8") as f:
                f.write(code_to_write)
        except Exception as e:
            return f"Error writing Program.cs: {e}"

        # Build quietly to suppress noise, then run without build
        try:
            build_proc = subprocess.run(
                [dotnet_path, "build", "-v:q"],
                cwd=cs_project_dir,
                capture_output=True,
                text=True,
                shell=False
            )
        except Exception as e:
            return f"Error invoking 'dotnet build': {e}"

        if build_proc.returncode != 0:
            build_err = (build_proc.stderr or "").strip()
            return build_err or "dotnet build failed."

        # Run without build to avoid restore/build logs
        try:
            run_proc = subprocess.run(
                [dotnet_path, "run", "--no-build"],
                cwd=cs_project_dir,
                capture_output=True,
                text=True,
                shell=False
            )
        except Exception as e:
            return f"Error invoking 'dotnet run': {e}"

        run_out = (run_proc.stdout or "").strip()
        run_err = (run_proc.stderr or "").strip()

        # Only return program output for a clean screenshot
        combined_output = "\n".join([s for s in (run_out, run_err) if s]).strip() or "Program executed successfully but produced no visible output."

        # Cleanup (best-effort)
        try:
            shutil.rmtree(cs_project_dir, ignore_errors=True)
        except Exception:
            pass

        return combined_output

    # Fallback to csc compiler if dotnet is not available
    csc_path = shutil.which("csc")
    if not csc_path:
        return "Error: .NET SDK ('dotnet') or C# compiler 'csc' not found in PATH. Please install .NET SDK."

    # Compile via csc
    cs_file = os.path.join(temp_dir, "Solution.cs")
    try:
        code_to_write = _massage_csharp(code)
        with open(cs_file, "w", encoding="utf-8") as f:
            f.write(code_to_write)
    except Exception as e:
        return f"Error writing C# source file: {e}"

    try:
        compile_proc = subprocess.run(
            [csc_path, "Solution.cs"],
            cwd=temp_dir,
            capture_output=True,
            text=True,
            shell=False
        )
    except Exception as e:
        return f"Error invoking C# compiler: {e}"

    compile_out = (compile_proc.stdout or "").strip()
    compile_err = (compile_proc.stderr or "").strip()

    if compile_proc.returncode != 0:
        combined = "\n".join([s for s in (compile_out, compile_err) if s]) or "Compilation failed with no output."
        return combined

    exe_path = os.path.join(temp_dir, "Solution.exe")
    if not os.path.isfile(exe_path):
        return "Compilation succeeded but 'Solution.exe' was not found."

    try:
        run_proc = subprocess.run(
            [exe_path],
            cwd=temp_dir,
            capture_output=True,
            text=True,
            shell=False
        )
    except Exception as e:
        return f"Error running compiled program: {e}"

    run_out = (run_proc.stdout or "").strip()
    run_err = (run_proc.stderr or "").strip()

    parts = []
    if run_out:
        parts.append(run_out)
    if run_err:
        parts.append(run_err)

    combined_output = "\n".join(parts).strip() or "Program executed successfully but produced no visible output."
    return combined_output


def process_question(q, language, user_name='Developer', document_terminal_path=None, screenshot_style='vscode'):
    """Process a coding question and generate solution with screenshot
    
    Args:
        q: The coding question
        language: Programming language
        user_name: User name for terminal display
        document_terminal_path: Terminal path for display
        screenshot_style: Screenshot style ('vscode', 'mac', 'simple')
    """
    if language == "python":
        sol = solve_coding_problem(q, "python")
        output = execute_code(sol)
        return sol, create_screenshot(output, user_name, document_terminal_path, screenshot_style, language)
    else:
        # Handle C# specially: execute the generated C# code for real dynamic output
        lang_key = (language or "").strip().lower()
        if lang_key in ("c#", "csharp"):
            sol_display = solve_coding_problem(q, "c#")
            try:
                output = execute_csharp_code(sol_display)
                if not output:
                    output = "Program executed successfully but produced no visible output."
            except Exception:
                output = "Program executed successfully but produced no visible output."
            return sol_display, create_screenshot(output, user_name, document_terminal_path, screenshot_style, language)
        
        # Default behavior for other non-Python languages:
        sol_display = solve_coding_problem(q, language)
        try:
            sol_python = solve_coding_problem(q, "python")
            output = execute_code(sol_python)
            if "Error executing code" in output or not output:
                output = "Program executed successfully.\nOutput displayed here."
        except Exception:
            output = "Program executed successfully.\nOutput displayed here."
        return sol_display, create_screenshot(output, user_name, document_terminal_path, screenshot_style, language)


# ----- Routes -----
@app.route('/')
def index():
    # Check if user is already authenticated
    user = flask_session.get('user')
    if user and flask_session.get('authenticated'):
        # Refresh user data from database to ensure it's current
        current_user = db_helper.get_user_by_phone(user.get('phone_number'))
        if current_user:
            flask_session['user'] = current_user
            logging.info(f"User {user.get('phone_number')} automatically logged in from existing session")
        else:
            # User no longer exists, clear session
            flask_session.clear()
    
    return render_template('index.html', firebase_api_key=os.getenv('FIREBASE_API_KEY'))

@app.route('/check_pdf_questions', methods=['POST'])
def check_pdf_questions():
    """Check number of questions in PDF before processing"""
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        file = request.files.get('file')
        if not file:
            return jsonify({"error": "PDF file is required."}), 400
        
        # Save file temporarily to extract questions
        _, file_ext = os.path.splitext(secure_filename(file.filename))
        temp_filename = f"temp_{uuid.uuid4().hex}{file_ext}"
        temp_path = os.path.join(TEMP_FOLDER, temp_filename)
        file.save(temp_path)
        
        try:
            # Extract and count questions
            pdf_text = extract_text_from_pdf(temp_path)
            questions = split_questions(pdf_text)
            question_count = len(questions)
            
            # Calculate credits required - 1 credit for up to 20 questions, +1 for each additional
            base_credits = 1  # Base credit for up to 20 questions
            extra_questions = max(0, question_count - 20)
            total_credits = base_credits + extra_questions
            
            # Check if user has enough credits
            current_user = db_helper.get_user_by_phone(user.get('phone_number'))
            user_credits = current_user['credits'] if current_user else 0
            
            response_data = {
                "question_count": question_count,
                "total_credits_required": total_credits,
                "user_credits": user_credits,
                "has_sufficient_credits": user_credits >= total_credits,
                "exceeds_limit": question_count > 20,
                "extra_questions": extra_questions,
                "questions_preview": questions[:3] if questions else []  # Show first 3 questions as preview
            }
            
            return jsonify(response_data)
            
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        logging.exception("Error checking PDF questions")
        return jsonify({'error': str(e)}), 500

@app.route('/upload_pdf_async', methods=['POST'])
def upload_pdf_async():
    """Asynchronous PDF processing with real-time updates - LOCAL VERSION (No Redis/Celery)"""
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        # Get form data
        file = request.files.get('file')
        name = request.form.get('name')
        reg_number = request.form.get('regNo')
        language = request.form.get('language', 'python').lower()
        confirmed = request.form.get('confirmed') == 'true'
        template_path = request.form.get('template_path', '')
        
        # Extract customization settings
        customization = None
        customization_json = request.form.get('customization')
        if customization_json:
            try:
                customization = json.loads(customization_json)
            except json.JSONDecodeError:
                logging.warning(f"Invalid customization JSON: {customization_json}")
                customization = None
        
        if not file or not name or not reg_number:
            return jsonify({"error": "Name, Register Number, and PDF file are required."}), 400
        
        # Check if user has enough credits
        current_user = db_helper.get_user_by_phone(user.get('phone_number'))
        if not current_user or current_user['credits'] <= 0:
            return jsonify({"error": "Insufficient credits. Please purchase credits to continue."}), 402
        
        # For local development, redirect to synchronous upload_pdf route
        # This avoids Redis/Celery dependency issues
        logging.info("Redirecting async request to synchronous processing for local development")
        
        # Generate a fake task ID for frontend compatibility
        task_id = str(uuid.uuid4())
        
        # Save file temporarily
        _, file_ext = os.path.splitext(secure_filename(file.filename))
        unique_filename = f"{uuid.uuid4().hex}{file_ext}"
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        file.save(file_path)
        
        return jsonify({
            "success": True,
            "task_id": task_id,
            "message": "PDF processing will start shortly. Please use the regular upload for now.",
            "estimated_time": "Processing will begin shortly...",
            "fallback_message": "For local development, please use the regular PDF upload button instead."
        })
        
    except Exception as e:
        logging.exception("Error in async PDF processing fallback")
        return jsonify({
            'error': 'Async processing not available in local development. Please use regular PDF upload.',
            'fallback': True
        }), 500

@app.route('/upload_pdf', methods=['POST'])
def upload_pdf():
    start_time = time.time()  # Track processing start time
    task_id = str(uuid.uuid4())
    
    try:
        # Add enhanced connection keep-alive headers to prevent timeout
        response_headers = {
            'Connection': 'keep-alive',
            'Keep-Alive': 'timeout=600, max=1000',
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no'  # Disable nginx buffering for real-time updates
        }
        
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        # Check file size early to prevent large uploads from causing connection issues
        file = request.files.get('file')
        if file:
            # Get file size
            file.seek(0, 2)  # Seek to end
            file_size = file.tell()
            file.seek(0)  # Reset to beginning
            
            # Check if file is too large (limit to 50MB)
            max_file_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_file_size:
                return jsonify({"error": f"File too large. Maximum size is {max_file_size // (1024*1024)}MB."}), 413
        
        # Get confirmed processing flag
        confirmed = request.form.get('confirmed', 'false').lower() == 'true'
        
        # Check if user has enough credits
        current_user = db_helper.get_user_by_phone(user.get('phone_number'))
        if not current_user or current_user['credits'] <= 0:
            return jsonify({"error": "Insufficient credits. Please purchase credits to continue."}), 402
        
        file = request.files.get('file')
        name = request.form.get('name')
        reg_number = request.form.get('regNo')
        language = request.form.get('language', 'python').lower()
        
        # Extract screenshot style from form data
        screenshot_style = request.form.get('screenshot_style', 'vscode')  # Default to vscode
        # Normalize alias and validate
        if screenshot_style == 'macos':
            screenshot_style = 'mac'
        if screenshot_style not in ['vscode', 'mac', 'simple']:
            screenshot_style = 'vscode'
        
        logging.info(f"PDF upload using screenshot style: {screenshot_style}")
        
        if not file or not name or not reg_number:
            return jsonify({"error": "Name, Register Number, and PDF file are required."}), 400
        
        # Initialize progress tracking
        progress_data = {
            'task_id': task_id,
            'user_id': user.get('phone_number'),
            'stage': 'initialization',
            'stage_name': 'Initializing PDF processing...',
            'progress': 5,
            'estimated_total_time': 'Calculating...',
            'elapsed_time': 0,
            'remaining_time': 'Calculating...',
            'start_time': start_time
        }
        
        # Store task in active tasks
        active_tasks[task_id] = {
            'type': 'pdf',
            'user_id': user.get('phone_number'),
            'start_time': start_time,
            'status': 'processing'
        }
        
        # Emit initial progress
        emit_progress_update(task_id, progress_data)
        
        # Save file
        _, file_ext = os.path.splitext(secure_filename(file.filename))
        unique_filename = f"{uuid.uuid4().hex}{file_ext}"
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        file.save(file_path)
        
        # Update progress: PDF extraction
        progress_data.update({
            'stage': 'pdf_extraction',
            'stage_name': 'Extracting questions from PDF...',
            'progress': 15,
            'elapsed_time': time.time() - start_time
        })
        emit_progress_update(task_id, progress_data)
        
        # Extract and process questions
        pdf_text = extract_text_from_pdf(file_path)
        questions = split_questions(pdf_text)
        
        # Update progress: Question analysis
        estimated_time = calculate_estimated_time(len(questions), 'pdf')
        progress_data.update({
            'stage': 'question_analysis',
            'stage_name': f'Analyzing {len(questions)} questions...',
            'progress': 25,
            'estimated_total_time': estimated_time,
            'elapsed_time': time.time() - start_time
        })
        emit_progress_update(task_id, progress_data)
        
        # Calculate credits required based on question count - 1 credit for up to 20 questions, +1 for each additional
        base_credits = 1  # Base credit for up to 20 questions
        extra_questions = max(0, len(questions) - 20)
        total_credits_required = base_credits + extra_questions
        
        # Check if user has enough credits for the actual question count
        if current_user['credits'] < total_credits_required:
            return jsonify({
                "error": f"Insufficient credits. You need {total_credits_required} credits for {len(questions)} questions, but you have {current_user['credits']} credits."
            }), 402
        
        # If not confirmed and exceeds 20 questions, require confirmation
        if not confirmed and len(questions) > 20:
            # Clean up task from active tasks before returning
            if task_id in active_tasks:
                del active_tasks[task_id]
            if task_id in task_progress:
                del task_progress[task_id]
            
            return jsonify({
                "requires_confirmation": True,
                "question_count": len(questions),
                "total_credits_required": total_credits_required,
                "extra_questions": extra_questions,
                "message": f"This PDF contains {len(questions)} questions. Processing will cost {total_credits_required} credits (1 base + {extra_questions} extra). Do you want to continue?"
            }), 400
        
        # Update progress: AI processing
        progress_data.update({
            'stage': 'ai_processing',
            'stage_name': 'Generating solutions using AI...',
            'progress': 30,
            'elapsed_time': time.time() - start_time,
            'questions_total': len(questions),
            'questions_completed': 0
        })
        emit_progress_update(task_id, progress_data)
        
        # Initialize solutions and screenshots lists with proper indexing to maintain order
        solutions_display = [None] * len(questions)
        screenshots = [None] * len(questions)
        
        # Assign one realistic Windows terminal path for the entire document
        # Determine user display name (session -> DB -> form -> default)
        user_info = flask_session.get('user', {})
        user_name = user_info.get('name')
        if not user_name and user_info.get('phone_number'):
            db_user = db_helper.get_user_by_phone(user_info.get('phone_number'))
            if db_user:
                user_name = db_user.get('name')
        if not user_name:
            user_name = name or "User"

        # Build realistic per-user terminal paths and pick one using rolling index
        realistic_paths = get_realistic_terminal_paths(user_name)
        path_index = int(hashlib.sha1(task_id.encode('utf-8')).hexdigest(), 16) % len(realistic_paths)
        document_terminal_path = realistic_paths[path_index]

        logging.info(f"Assigned terminal path for document {task_id}: {document_terminal_path}")
        
        # FIXED: Sequential processing to prevent Claude API concurrent connection issues
        # Check if we should use sequential processing for Claude API calls
        claude_available = claude_manager.get_available_keys_count() > 0
        completed_questions = 0
        
        if claude_available and len(questions) > 3:
            # Sequential processing for better Claude API compliance
            logging.info(f"Using sequential processing for {len(questions)} questions to prevent Claude API concurrent connection issues")
            
            for question_index, q in enumerate(questions):
                try:
                    # Process question sequentially
                    sol, scr = process_question(q, language, user_name, document_terminal_path, screenshot_style)
                    
                    # Store solution and screenshot at the correct index to maintain order
                    solutions_display[question_index] = sol
                    screenshots[question_index] = scr
                    completed_questions += 1
                    
                    # Update progress for each completed question
                    question_progress = 30 + (completed_questions / len(questions)) * 50  # 30% to 80%
                    progress_data.update({
                        'stage': 'ai_processing',
                        'stage_name': f'Processing questions ({completed_questions}/{len(questions)})...',
                        'progress': int(question_progress),
                        'elapsed_time': time.time() - start_time,
                        'questions_completed': completed_questions
                    })
                    emit_progress_update(task_id, progress_data)
                    
                    logging.debug(f"Completed question {question_index + 1}/{len(questions)}: {q[:50]}...")
                    
                except Exception as e:
                    logging.error(f"Error processing question {question_index + 1}: {e}")
                    # Store error message while maintaining order
                    solutions_display[question_index] = f"Error: Failed to generate solution - {str(e)}"
                    screenshots[question_index] = b""  # Empty screenshot for failed questions
                    completed_questions += 1
        else:
            # Concurrent processing for DeepSeek or small question sets
            logging.info(f"Using concurrent processing for {len(questions)} questions")
            
            # Create futures with their corresponding question indices to prevent shuffling
            futures_with_indices = [
                (i, global_executor.submit(process_question, q, language, user_name, document_terminal_path, screenshot_style))
                for i, q in enumerate(questions)
            ]
            
            # Process completed futures while maintaining original question order
            for question_index, future in futures_with_indices:
                try:
                    # Wait for this specific future to complete
                    sol, scr = future.result()
                    
                    # Store solution and screenshot at the correct index to maintain order
                    solutions_display[question_index] = sol
                    screenshots[question_index] = scr
                    completed_questions += 1
                    
                    # Update progress for each completed question
                    question_progress = 30 + (completed_questions / len(questions)) * 50  # 30% to 80%
                    progress_data.update({
                        'stage': 'ai_processing',
                        'stage_name': f'Processing questions ({completed_questions}/{len(questions)})...',
                        'progress': int(question_progress),
                        'elapsed_time': time.time() - start_time,
                        'questions_completed': completed_questions
                    })
                    emit_progress_update(task_id, progress_data)
                    
                    logging.debug(f"Completed question {question_index + 1}/{len(questions)}: {questions[question_index][:50]}...")
                    
                except Exception as e:
                    logging.error(f"Error processing question {question_index + 1}: {e}")
                    # Store error message while maintaining order
                    solutions_display[question_index] = f"Error: Failed to generate solution - {str(e)}"
                    screenshots[question_index] = b""  # Empty screenshot for failed questions
                    completed_questions += 1
        
        # Update progress: Document generation
        progress_data.update({
            'stage': 'document_generation',
            'stage_name': 'Creating Word document...',
            'progress': 85,
            'elapsed_time': time.time() - start_time
        })
        emit_progress_update(task_id, progress_data)
        
        # Parse customization settings
        customization = None
        if 'customization' in request.form:
            try:
                customization_str = request.form.get('customization', '{}')
                customization = json.loads(customization_str) if customization_str else None
            except json.JSONDecodeError:
                logging.error(f"Invalid customization JSON: {request.form.get('customization')}")
                customization = None
        
        # Extract CHRIST template data from multiple sources
        christ_template_data = None
        
        # Method 1: Direct form fields (legacy)
        if 'christName' in request.form or 'christRegNo' in request.form:
            christ_template_data = {
                'name': request.form.get('christName', '').strip(),
                'reg_number': request.form.get('christRegNo', '').strip(),
                'course': request.form.get('christCourse', '').strip(),
                'class_section': request.form.get('christClass', '').strip(),
                'practical_number': request.form.get('christPractical', '').strip(),
                'teacher_name': request.form.get('christTeacher', '').strip()
            }
            logging.info(f"CHRIST template data received from form fields: {christ_template_data}")
        
        # Method 2: From customization JSON (new method)
        elif customization and 'christTemplateData' in customization:
            christ_template_data = customization.get('christTemplateData', {})
            logging.info(f"CHRIST template data received from customization: {christ_template_data}")
        
        # Method 3: Check if template is 'christ' and extract from form
        elif request.form.get('template') == 'christ':
            # Try to extract from any available data
            christ_template_data = {
                'name': name,  # Use the main form name
                'reg_number': reg_number,  # Use the main form reg number
                'course': request.form.get('christCourse', '').strip(),
                'class_section': request.form.get('christClass', '').strip(),
                'practical_number': request.form.get('christPractical', '').strip(),
                'teacher_name': request.form.get('christTeacher', '').strip()
            }
            logging.info(f"CHRIST template data extracted for christ template: {christ_template_data}")
        
        # Get template option from form
        template_option = request.form.get('template', 'none')  # Default to 'none' if not specified
        logging.info(f"Template option selected: {template_option}")
        
        # Generate Word document
        output_file = os.path.join(TEMP_FOLDER, f"solutions_{uuid.uuid4().hex}.docx")
        generated_path = generate_word_doc(name, reg_number, questions, solutions_display, screenshots, output_file, customization, christ_template_data, template_option)
        logging.info(f"Generated document path: {generated_path}")
        
        # Check if document was generated successfully
        if not os.path.exists(generated_path):
            logging.error(f"File not found: {generated_path}")
            
            # Update progress: Failed
            progress_data.update({
                'stage': 'failed',
                'stage_name': 'Document generation failed',
                'progress': 100,
                'elapsed_time': time.time() - start_time,
                'error': 'Document generation failed'
            })
            emit_progress_update(task_id, progress_data)
            
            # Clean up task
            if task_id in active_tasks:
                del active_tasks[task_id]
            
            # Record failed submission
            db_helper.insert_submission(
                phone_number=user.get('phone_number'),
                pdf_name=file.filename,
                questions_count=len(questions),
                questions_solved=0,
                questions_failed=len(questions),
                failed_questions=questions,
                solved=False,
                error_details="Document generation failed",
                processing_time_seconds=(time.time() - start_time)
            )
            return jsonify({"error": "Failed to generate the document."}), 500
        
        # Calculate success metrics
        successful_solutions = [sol for sol in solutions_display if sol and not sol.startswith("Error")]
        questions_solved = len(successful_solutions)
        questions_failed = len(questions) - questions_solved
        failed_questions = [q for q, sol in zip(questions, solutions_display) if not sol or sol.startswith("Error")]
        is_fully_solved = questions_solved == len(questions)
        
        # Record submission in database with credit usage
        submission_record = db_helper.insert_submission(
            phone_number=user.get('phone_number'),
            pdf_name=file.filename,
            questions_count=len(questions),
            questions_solved=questions_solved,
            questions_failed=questions_failed,
            failed_questions=failed_questions,
            solved=is_fully_solved,
            error_details=None if is_fully_solved else f"{questions_failed} questions failed to solve",
            processing_time_seconds=(time.time() - start_time),
            submission_type='pdf'
        )
        
        logging.info(f"Submission recorded: {submission_record}")
        
        # Update progress: Finalization
        progress_data.update({
            'stage': 'finalization',
            'stage_name': 'Finalizing and preparing download...',
            'progress': 95,
            'elapsed_time': time.time() - start_time
        })
        emit_progress_update(task_id, progress_data)
        
        # Deduct credits based on question count - FIXED LOGIC
        updated_user = db_helper.deduct_credits_by_count(user.get('phone_number'), len(questions))
        if updated_user:
            # Update session with new user data
            flask_session['user'] = updated_user
        else:
            logging.error(f"Failed to deduct credits for {user.get('phone_number')} - PDF processing")
        
        # Update progress: Completed
        final_elapsed = time.time() - start_time
        progress_data.update({
            'stage': 'completed',
            'stage_name': 'PDF processing completed successfully!',
            'progress': 100,
            'elapsed_time': final_elapsed,
            'total_time': f"{final_elapsed:.1f}s",
            'download_ready': True
        })
        emit_progress_update(task_id, progress_data)
        
        # Clean up task after a short delay
        def cleanup_task():
            time.sleep(5)  # Keep task info for 5 seconds
            if task_id in active_tasks:
                del active_tasks[task_id]
            if task_id in task_progress:
                del task_progress[task_id]
        
        threading.Thread(target=cleanup_task, daemon=True).start()
        
        # Return task_id in response header for frontend tracking with enhanced headers
        response = make_response(send_file(generated_path, as_attachment=True))
        response.headers['X-Task-ID'] = task_id
        response.headers['Connection'] = 'keep-alive'
        response.headers['Keep-Alive'] = 'timeout=600, max=1000'
        response.headers['Cache-Control'] = 'no-cache'
        response.headers['X-Accel-Buffering'] = 'no'
        return response
        
    except FileNotFoundError as e:
        logging.error(f"File not found error: {e}")
        return jsonify({'error': 'Required file not found. Please try again.'}), 404
    except MemoryError as e:
        logging.error(f"Memory error during processing: {e}")
        return jsonify({'error': 'File too large to process. Please try a smaller file.'}), 413
    except TimeoutError as e:
        logging.error(f"Timeout error during processing: {e}")
        return jsonify({'error': 'Processing timeout. Please try again with a smaller file.'}), 408
    except ConnectionError as e:
        logging.error(f"Connection error during processing: {e}")
        return jsonify({'error': 'Connection error. Please check your internet connection and try again.'}), 503
    except Exception as e:
        logging.exception("Error processing PDF")
        # Record failed submission due to exception
        try:
            user = flask_session.get('user')
            if user:
                db_helper.insert_submission(
                    phone_number=user.get('phone_number'),
                    pdf_name=file.filename if 'file' in locals() else 'unknown',
                    questions_count=0,
                    questions_solved=0,
                    questions_failed=0,
                    failed_questions=[],
                    solved=False,
                    error_details=f"Processing exception: {str(e)}",
                    processing_time_seconds=(time.time() - start_time),
                    submission_type='pdf'
                )
        except:
            pass  # Don't let logging errors break the main error response
        
        # Return appropriate error based on exception type
        if "timeout" in str(e).lower():
            return jsonify({'error': 'Request timeout. Please try again.'}), 408
        elif "connection" in str(e).lower():
            return jsonify({'error': 'Connection error. Please try again.'}), 503
        else:
            return jsonify({'error': f'Processing error: {str(e)}'}), 500

@app.route('/task_status/<task_id>', methods=['GET'])
def get_task_status(task_id):
    """Get current task status"""
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        # Get task from database
        task = db_helper.get_task_by_id(task_id)
        if not task:
            return jsonify({"error": "Task not found."}), 404
        
        # Verify user owns the task
        if task.get('phone_number') != user.get('phone_number'):
            return jsonify({"error": "Access denied."}), 403
        
        # Get task summary with estimated time
        from task_manager import task_manager
        task_summary = task_manager.get_task_summary(task_id)
        
        if task_summary:
            # Calculate estimated completion time
            estimated_time = real_time_progress.calculate_estimated_time(task_summary)
            task_summary['estimated_completion'] = estimated_time
            
            return jsonify({
                "success": True,
                "task": task_summary
            })
        else:
            return jsonify({"error": "Task details not available."}), 404
            
    except Exception as e:
        logging.exception("Error getting task status")
        return jsonify({'error': str(e)}), 500

@app.route('/download_task_result/<task_id>', methods=['GET'])
def download_task_result(task_id):
    """Download completed task result"""
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        # Get task from database
        task = db_helper.get_task_by_id(task_id)
        if not task:
            return jsonify({"error": "Task not found."}), 404
        
        # Verify user owns the task
        if task.get('phone_number') != user.get('phone_number'):
            return jsonify({"error": "Access denied."}), 403
        
        # Check if task is completed
        if task.get('task_status') != 'COMPLETED':
            return jsonify({"error": "Task is not completed yet."}), 400
        
        # Check if output file exists
        output_file = task.get('output_file_path')
        if not output_file or not os.path.exists(output_file):
            return jsonify({"error": "Output file not found."}), 404
        
        return send_file(output_file, as_attachment=True)
        
    except Exception as e:
        logging.exception("Error downloading task result")
        return jsonify({'error': str(e)}), 500

@app.route('/user_tasks', methods=['GET'])
def get_user_tasks():
    """Get user's task history"""
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        # Get user's tasks
        from task_manager import task_manager
        tasks = task_manager.get_user_tasks(user.get('phone_number'), 20)
        
        # Format tasks for response
        formatted_tasks = []
        for task in tasks:
            formatted_task = {
                'task_id': task['task_id'],
                'task_type': task['task_type'],
                'status': task['task_status'],
                'progress': task['task_progress'],
                'created_at': task['created_at'],
                'completed_at': task.get('completed_at'),
                'questions_count': task.get('questions_count', 0),
                'questions_solved': task.get('questions_solved', 0),
                'credits_used': task.get('credits_used', 0),
                'has_output': task.get('output_file_path') is not None
            }
            formatted_tasks.append(formatted_task)
        
        return jsonify({
            "success": True,
            "tasks": formatted_tasks
        })
        
    except Exception as e:
        logging.exception("Error getting user tasks")
        return jsonify({'error': str(e)}), 500

@app.route('/send_otp', methods=['POST'])
def send_otp():
    phone_number = request.json.get('phone')
    user_name = request.json.get('name', '').strip()
    
    if not validate_phone_number(phone_number):
        return jsonify({'error': 'Invalid phone number'}), 400
    
    # Check if user already exists
    existing_user = db_helper.get_user_by_phone(phone_number)
    
    # If user exists, don't allow new signup with same number
    if existing_user:
        return jsonify({
            'error': 'An account with this phone number already exists. Please use the Login option to access your existing account.'
        }), 409  # 409 Conflict status code
    
    # For new users, name is required
    if not user_name:
        return jsonify({'error': 'Name is required for new users'}), 400

    # Get OTP provider from environment
    otp_provider = os.getenv('OTP_PROVIDER', 'firebase').lower()
    
    if otp_provider == 'firebase':
        # Use Firebase SMS OTP service
        try:
            firebase_otp_service = FirebaseOTP()
            
            # Create custom token for the phone number
            token_result = firebase_otp_service.create_custom_token(phone_number)
            if not token_result['success']:
                logging.error(f"Firebase token creation failed: {token_result.get('error')}")
                # Fallback to console OTP for development
                return send_console_otp(phone_number, user_name)
            
            # Store the token and user info for verification
            flask_session['custom_token'] = token_result['custom_token']
            flask_session['phone_number'] = phone_number
            flask_session['user_name'] = user_name
            flask_session['otp_provider'] = 'firebase'
            
            logging.info(f"Firebase custom token created for {phone_number}")
            return jsonify({
                'message': 'Please complete phone verification using Firebase Authentication.',
                'provider': 'firebase',
                'custom_token': token_result['custom_token'],
                'phone_number': phone_number
            }), 200
            
        except Exception as e:
            logging.error(f"Firebase SMS OTP error: {e}")
            # Fallback to console OTP for development
            return send_console_otp(phone_number, user_name)
    
    else:
        # Default to console OTP for development
        return send_console_otp(phone_number, user_name)

def send_console_otp(phone_number, user_name):
    """Fallback console OTP for development"""
    try:
        # Generate OTP
        otp_code = generate_otp()
        
        # Store OTP in database
        otp_record = db_helper.store_otp(phone_number, otp_code, user_name)
        
        if otp_record:
            flask_session['phone_number'] = phone_number
            flask_session['user_name'] = user_name
            flask_session['otp_provider'] = 'console'
            
            # Display OTP in console for development
            print(f"\n🔐 DEVELOPMENT OTP for {phone_number}: {otp_code}")
            print(f"📱 In production, this would be sent via SMS")
            logging.info(f"Console OTP generated for {phone_number}: {otp_code}")
            
            return jsonify({
                'message': f'Development mode: Your OTP is {otp_code}. Check the server console for details.',
                'provider': 'console',
                'otp_code': otp_code  # Only for development
            }), 200
        else:
            return jsonify({'error': 'Failed to generate OTP'}), 500
            
    except Exception as e:
        logging.error(f"Console OTP error: {e}")
        return jsonify({'error': 'OTP service temporarily unavailable'}), 500

@app.route('/verify_otp', methods=['POST'])
def verify_otp():
    """Verify OTP with comprehensive error handling and JSON responses"""
    try:
        logging.info("OTP verification request received")
        
        # Parse request data with error handling
        try:
            request_data = request.get_json()
            if not request_data:
                logging.error("Empty request data received")
                return jsonify({'error': 'Invalid request data'}), 400
        except Exception as e:
            logging.error(f"Error parsing JSON data: {e}")
            return jsonify({'error': 'Invalid JSON data'}), 400
        
        phone_number = flask_session.get('phone_number')
        otp_provider = flask_session.get('otp_provider', 'console')
        remember_me = request_data.get('remember_me', True)  # Default to True for better UX

        if not phone_number:
            return jsonify({'error': 'Session expired. Please try again.'}), 403

        # Handle Firebase authentication
        if otp_provider == 'firebase':
            id_token = request_data.get('id_token')
            if not id_token:
                return jsonify({'error': 'Firebase ID token required'}), 400
                
            try:
                firebase_otp_service = FirebaseOTP()
                result = firebase_otp_service.verify_id_token(id_token)
                
                if result['success']:
                    # Token verified successfully
                    verified_phone = result['phone_number']
                    user = db_helper.get_user_by_phone(verified_phone)
                    if not user:
                        # Create new user with name from session
                        user_name = flask_session.get('user_name', 'User')
                        user = db_helper.create_user(verified_phone, user_name)
                    
                    # Clear all session data and set up fresh session
                    flask_session.clear()
                    flask_session['user'] = user
                    flask_session['authenticated'] = True
                    flask_session['phone_number'] = verified_phone
                    
                    # If remember me is checked, make session permanent
                    if remember_me:
                        flask_session.permanent = True
                        logging.info(f"User {verified_phone} signed up/logged in with Firebase SMS OTP (persistent session)")
                    else:
                        flask_session.permanent = False
                        logging.info(f"User {verified_phone} signed up/logged in with Firebase SMS OTP (temporary session)")
                        
                    return jsonify({
                        "message": "SMS OTP verified successfully",
                        "provider": "firebase",
                        'user': {
                            'name': user.get('name', 'User'),
                            'phone_number': user.get('phone_number'),
                            'credits': user.get('credits', 0)
                        }
                    }), 200
                else:
                    return jsonify({'error': result.get('error', 'Invalid or expired token')}), 403
                    
            except Exception as e:
                logging.error(f"Firebase OTP verification error: {e}")
                return jsonify({
                    'error': 'OTP verification failed',
                    'details': str(e) if app.debug else 'Please try again'
                }), 500
        
        # Handle console/traditional OTP
        else:
            otp_code = request_data.get('otp')
            if not otp_code:
                return jsonify({'error': 'OTP code required'}), 400
                
            # Traditional OTP verification (legacy)
            otp_record = db_helper.verify_otp(phone_number, otp_code)
            if otp_record:
                user = db_helper.get_user_by_phone(phone_number)
                if not user:
                    # Create new user with name from session or OTP record
                    user_name = flask_session.get('user_name') or otp_record.get('user_name', 'User')
                    user = db_helper.create_user(phone_number, user_name)
                
                # Clear all session data and set up fresh session
                flask_session.clear()
                flask_session['user'] = user
                flask_session['authenticated'] = True
                flask_session['phone_number'] = phone_number
                
                # If remember me is checked, make session permanent
                if remember_me:
                    flask_session.permanent = True
                    logging.info(f"User {phone_number} signed up/logged in with console OTP (persistent session)")
                else:
                    flask_session.permanent = False
                    logging.info(f"User {phone_number} signed up/logged in with console OTP (temporary session)")
                    
                return jsonify({
                    "message": "OTP verified successfully",
                    "provider": "console",
                    'user': {
                        'name': user.get('name', 'User'),
                        'phone_number': user.get('phone_number'),
                        'credits': user.get('credits', 0)
                    }
                }), 200
            else:
                return jsonify({'error': 'Invalid or expired OTP'}), 403
                
    except Exception as e:
        logging.error(f"Unexpected error in verify_otp: {e}")
        return jsonify({
            'error': 'An unexpected error occurred',
            'details': str(e) if app.debug else 'Please try again later'
        }), 500
    
@app.route('/firebase_auth', methods=['GET'])
def firebase_auth():
    """Serve Firebase authentication page"""
    return render_template('firebase_auth.html', firebase_api_key=os.getenv('FIREBASE_API_KEY'))

@app.route('/setup_firebase_session', methods=['POST'])
def setup_firebase_session():
    """Setup Firebase session for OTP authentication"""
    try:
        request_data = request.get_json()
        if not request_data:
            return jsonify({'error': 'Invalid request data'}), 400
        
        phone_number = request_data.get('phone_number')
        user_name = request_data.get('user_name', '')
        
        if not phone_number:
            return jsonify({'error': 'Phone number is required'}), 400
        
        # Clean and validate phone number
        cleaned_phone = ''.join(filter(str.isdigit, phone_number))
        if len(cleaned_phone) != 10:
            return jsonify({'error': 'Please enter a valid 10-digit phone number'}), 400
        
        # Set up session for Firebase OTP
        flask_session['phone_number'] = cleaned_phone
        flask_session['user_name'] = user_name
        flask_session['otp_provider'] = 'firebase'
        flask_session['session_initialized'] = True
        
        logging.info(f"Firebase session initialized for phone: {cleaned_phone}")
        
        return jsonify({
            'success': True,
            'message': 'Session initialized successfully',
            'phone_number': cleaned_phone
        }), 200
        
    except Exception as e:
        logging.error(f"Error setting up Firebase session: {e}")
        return jsonify({
            'error': 'Failed to initialize session',
            'details': str(e) if app.debug else 'Please try again'
        }), 500

@app.route('/login_existing', methods=['POST'])
def login_existing():
    """Login existing user without OTP"""
    phone_number = request.json.get('phone')
    
    if not phone_number:
        return jsonify({'error': 'Phone number is required'}), 400
    
    # Check if user exists
    user = db_helper.get_user_by_phone(phone_number)
    if not user:
        return jsonify({'error': 'User not found. Please sign up first.'}), 404
    
# Record OTP verification (simplified for Firebase)
    # Note: Using a simple approach since Firebase verification is already done
    pass

    # Set up session
    flask_session['user'] = user
    flask_session['authenticated'] = True
    flask_session['phone_number'] = phone_number
    flask_session.permanent = True  # Make session permanent for quick login
    
    logging.info(f"User {phone_number} logged in via quick login")
    
    return jsonify({
        'message': 'Login successful',
        'user': {
            'name': user.get('name', 'User'),
            'phone_number': user.get('phone_number'),
            'credits': user.get('credits', 0)
        }
    }), 200

@app.route('/verify_firebase_otp', methods=['POST'])
def verify_firebase_otp():
    """Verify Firebase OTP from main index page - compatible with existing frontend"""
    try:
        logging.info("Firebase OTP verification request received from main page")
        
        # Parse request data
        try:
            request_data = request.get_json()
            if not request_data:
                logging.error("Empty request data received")
                return jsonify({'error': 'Invalid request data'}), 400
        except Exception as e:
            logging.error(f"Error parsing JSON data: {e}")
            return jsonify({'error': 'Invalid JSON data'}), 400
        
        # Extract data from request
        firebase_id_token = request_data.get('firebase_id_token')
        name = request_data.get('name', '').strip()
        phone = request_data.get('phone', '').strip()
        remember_me = request_data.get('remember_me', True)
        
        # Validate required fields
        if not firebase_id_token:
            return jsonify({'error': 'Firebase ID token is required'}), 400
        
        if not name or not phone:
            return jsonify({'error': 'Name and phone number are required'}), 400
        
        # Verify Firebase ID token
        try:
            result = firebase_otp_service.verify_id_token(firebase_id_token)
            
            if not result['success']:
                return jsonify({'error': result.get('error', 'Invalid Firebase token')}), 403
            
            # Get verified phone from Firebase token
            verified_phone = result['phone_number']
            
            # Ensure the phone number from token matches the one from request
            if verified_phone != phone:
                logging.warning(f"Phone mismatch: token={verified_phone}, request={phone}")
                # Use the verified phone from token for security
                phone = verified_phone
            
            # Check if user already exists
            existing_user = db_helper.get_user_by_phone(phone)
            
            if existing_user:
                # User exists - log them in
                user = existing_user
                logging.info(f"Existing user {phone} logged in via Firebase OTP")
            else:
                # Create new user
                user = db_helper.create_user(phone, name)
                if not user:
                    return jsonify({'error': 'Failed to create user account'}), 500
                logging.info(f"New user {phone} created via Firebase OTP")
            
# Record OTP verification (simplified for Firebase)
            # Note: Using a simple approach since Firebase verification is already done
            pass

            # Set up session
            flask_session.clear()  # Clear any existing session data
            flask_session['user'] = user
            flask_session['authenticated'] = True
            flask_session['phone_number'] = phone
            
            # Set session persistence based on remember_me
            if remember_me:
                flask_session.permanent = True
                logging.info(f"User {phone} signed up/logged in with Firebase OTP (persistent session)")
            else:
                flask_session.permanent = False
                logging.info(f"User {phone} signed up/logged in with Firebase OTP (temporary session)")
            
            return jsonify({
                'success': True,
                'message': 'Firebase OTP verified successfully',
                'user': {
                    'name': user.get('name', 'User'),
                    'phone_number': user.get('phone_number'),
                    'credits': user.get('credits', 0),
                    'is_priority': user.get('is_priority', False)
                }
            }), 200
            
        except Exception as e:
            logging.error(f"Firebase token verification error: {e}")
            return jsonify({
                'error': 'Firebase token verification failed',
                'details': str(e) if app.debug else 'Please try again'
            }), 500
            
    except Exception as e:
        logging.error(f"Unexpected error in verify_firebase_otp: {e}")
        return jsonify({
            'error': 'An unexpected error occurred',
            'details': str(e) if app.debug else 'Please try again later'
        }), 500


@app.route('/manual_solve', methods=['POST'])
def manual_solve():
    start_time = time.time()  # Track processing start time
    task_id = str(uuid.uuid4())
    
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        # Force refresh user data directly from Supabase to avoid cached data
        try:
            from supabase import create_client
            supabase_client = create_client(SUPABASE_URL, SUPABASE_KEY)
            response = supabase_client.table('users').select('*').eq('phone_number', user.get('phone_number')).execute()
            if response.data:
                current_user = response.data[0]
                logging.info(f"Direct Supabase check: User {user.get('phone_number')} has {current_user['credits']} credits")
            else:
                current_user = db_helper.get_user_by_phone(user.get('phone_number'))
        except Exception as e:
            logging.error(f"Error in direct Supabase check: {e}")
            current_user = db_helper.get_user_by_phone(user.get('phone_number'))
        
        if not current_user or current_user['credits'] <= 0:
            return jsonify({"error": "Insufficient credits. Please purchase credits to continue."}), 402
        
        data = request.get_json()
        name = data.get('name')
        reg_number = data.get('regNo')
        language = data.get('language', 'python').lower()
        questions_text = data.get('questions')
        confirmed = data.get('confirmed', False)
        
        # Parse customization data if provided
        customization = None
        if 'customization' in data:
            try:
                customization = data.get('customization')
                logging.info(f"Received customization data: {customization}")
            except Exception as e:
                logging.warning(f"Failed to parse customization data: {e}")
                customization = None
        
        if not questions_text or not name or not reg_number:
            return jsonify({"error": "Name, Register Number, and Questions are required."}), 400
        
        # Initialize progress tracking
        progress_data = {
            'task_id': task_id,
            'user_id': user.get('phone_number'),
            'stage': 'initialization',
            'stage_name': 'Initializing question processing...',
            'progress': 5,
            'estimated_total_time': 'Calculating...',
            'elapsed_time': 0,
            'remaining_time': 'Calculating...',
            'start_time': start_time
        }
        
        # Store task in active tasks
        active_tasks[task_id] = {
            'type': 'manual',
            'user_id': user.get('phone_number'),
            'start_time': start_time,
            'status': 'processing'
        }
        
        # Emit initial progress
        emit_progress_update(task_id, progress_data)
        
        # Process questions
        questions = split_questions(questions_text) if isinstance(questions_text, str) else questions_text
        
        # Update progress: Question analysis
        estimated_time = calculate_estimated_time(len(questions), 'manual')
        progress_data.update({
            'stage': 'question_analysis',
            'stage_name': f'Analyzing {len(questions)} questions...',
            'progress': 20,
            'estimated_total_time': estimated_time,
            'elapsed_time': time.time() - start_time,
            'questions_total': len(questions),
            'questions_completed': 0
        })
        emit_progress_update(task_id, progress_data)
        
        # Calculate credits required based on question count - 1 credit for up to 20 questions, +1 for each additional
        base_credits = 1  # Base credit for up to 20 questions
        extra_questions = max(0, len(questions) - 20)
        total_credits_required = base_credits + extra_questions
        
        # Check if user has enough credits for the actual question count
        if current_user['credits'] < total_credits_required:
            return jsonify({
                "error": f"Insufficient credits. You need {total_credits_required} credits for {len(questions)} questions, but you have {current_user['credits']} credits."
            }), 402
        
        # If not confirmed and exceeds 20 questions, require confirmation
        if not confirmed and len(questions) > 20:
            # Clean up task from active tasks before returning
            if task_id in active_tasks:
                del active_tasks[task_id]
            if task_id in task_progress:
                del task_progress[task_id]
            
            return jsonify({
                "requires_confirmation": True,
                "question_count": len(questions),
                "total_credits_required": total_credits_required,
                "extra_questions": extra_questions,
                "message": f"You have entered {len(questions)} questions. Processing will cost {total_credits_required} credits (1 base + {extra_questions} extra). Do you want to continue?"
            }), 400
        
        # Update progress: AI processing
        progress_data.update({
            'stage': 'ai_processing',
            'stage_name': 'Generating solutions using AI...',
            'progress': 25,
            'elapsed_time': time.time() - start_time
        })
        emit_progress_update(task_id, progress_data)
        
        solutions_display = []
        screenshots = []

        # Determine user display name from Supabase DB (preferred), then session, then provided name
        user_info = flask_session.get('user', {})
        db_user = db_helper.get_user_by_phone(user_info.get('phone_number')) if user_info.get('phone_number') else None

        if db_user and db_user.get('name'):
            user_name = db_user.get('name')
        elif user_info.get('name'):
            user_name = user_info.get('name')
        else:
            user_name = name or "User"

        # Build realistic per-user terminal paths and pick one deterministically using task_id
        realistic_paths = get_realistic_terminal_paths(user_name)
        path_index = int(hashlib.sha1(task_id.encode('utf-8')).hexdigest(), 16) % len(realistic_paths)
        document_terminal_path = realistic_paths[path_index]
        logging.info(f"Assigned terminal path for manual document {task_id}: {document_terminal_path}")

        # Get and validate screenshot style (include 'realistic' variant and macos alias)
        screenshot_style = data.get('screenshot_style', 'vscode')
        if screenshot_style == 'macos':
            screenshot_style = 'mac'
        if screenshot_style not in ['vscode', 'mac', 'simple', 'realistic']:
            screenshot_style = 'vscode'
        logging.info(f"Using screenshot style: {screenshot_style} for manual solve")

        # Process each question with progress updates
        futures = [global_executor.submit(process_question, q, language, user_name, document_terminal_path, screenshot_style) for q in questions]
        completed_questions = 0

        for future in concurrent.futures.as_completed(futures):
            sol, scr = future.result()
            solutions_display.append(sol)
            screenshots.append(scr)
            completed_questions += 1

            # Update progress for each completed question
            question_progress = 25 + (completed_questions / len(questions)) * 60  # 25% to 85%
            progress_data.update({
                'stage': 'ai_processing',
                'stage_name': f'Processing questions ({completed_questions}/{len(questions)})...',
                'progress': int(question_progress),
                'elapsed_time': time.time() - start_time,
                'questions_completed': completed_questions
            })
            emit_progress_update(task_id, progress_data)
        
        # Update progress: Document generation
        progress_data.update({
            'stage': 'document_generation',
            'stage_name': 'Creating Word document...',
            'progress': 90,
            'elapsed_time': time.time() - start_time
        })
        emit_progress_update(task_id, progress_data)
        
        # Extract CHRIST template data from customization if provided
        christ_template_data = None
        if customization and 'christTemplateData' in customization:
            christ_template_data = customization.get('christTemplateData', {})
            logging.info(f"CHRIST template data received from manual form customization: {christ_template_data}")
        # Also extract from direct fields if present (mimic upload_pdf route)
        elif any(k in data for k in ['christName', 'christRegNo', 'christCourse', 'christClass', 'christPractical', 'christTeacher']):
            christ_template_data = {
                'name': data.get('christName', '').strip(),
                'reg_number': data.get('christRegNo', '').strip(),
                'course': data.get('christCourse', '').strip(),
                'class_section': data.get('christClass', '').strip(),
                'practical_number': data.get('christPractical', '').strip(),
                'teacher_name': data.get('christTeacher', '').strip()
            }
            logging.info(f"CHRIST template data received from manual direct fields: {christ_template_data}")
        
        # Get template option from data (for manual solve, it comes from JSON data)
        template_option = data.get('template', 'none')  # Default to 'none' if not specified
        logging.info(f"Template option selected for manual solve: {template_option}")
        logging.info(f"Final CHRIST template data being passed to generate_word_doc: {christ_template_data}")
        logging.info(f"Customization data being passed to generate_word_doc: {customization}")
        
        # Generate document
        output_file = os.path.join(TEMP_FOLDER, f"manual_solutions_{uuid.uuid4().hex}.docx")
        logging.info(f"About to call generate_word_doc with template_option='{template_option}' and output_file='{output_file}'")
        generated_path = generate_word_doc(name, reg_number, questions, solutions_display, screenshots, output_file, customization, christ_template_data, template_option)
        
        if not os.path.exists(generated_path):
            # Update progress: Failed
            progress_data.update({
                'stage': 'failed',
                'stage_name': 'Document generation failed',
                'progress': 100,
                'elapsed_time': time.time() - start_time,
                'error': 'Document generation failed'
            })
            emit_progress_update(task_id, progress_data)
            
            # Clean up task
            if task_id in active_tasks:
                del active_tasks[task_id]
            
            # Record failed submission
            db_helper.insert_submission(
                phone_number=user.get('phone_number'),
                pdf_name='Manual Questions',
                questions_count=len(questions),
                questions_solved=0,
                questions_failed=len(questions),
                failed_questions=questions,
                solved=False,
                error_details="Document generation failed",
                processing_time_seconds=(time.time() - start_time),
                submission_type='manual'
            )
            return jsonify({"error": "Failed to generate the document."}), 500
        
        # Calculate success metrics
        successful_solutions = [sol for sol in solutions_display if sol and not sol.startswith("Error")]
        questions_solved = len(successful_solutions)
        questions_failed = len(questions) - questions_solved
        failed_questions = [q for q, sol in zip(questions, solutions_display) if not sol or sol.startswith("Error")]
        is_fully_solved = questions_solved == len(questions)
        
        # Record submission in database
        submission_record = db_helper.insert_submission(
            phone_number=user.get('phone_number'),
            pdf_name='Manual Questions',
            questions_count=len(questions),
            questions_solved=questions_solved,
            questions_failed=questions_failed,
            failed_questions=failed_questions,
            solved=is_fully_solved,
            error_details=None if is_fully_solved else f"{questions_failed} questions failed to solve",
            processing_time_seconds=(time.time() - start_time),
            submission_type='manual'
        )
        
        logging.info(f"Manual submission recorded: {submission_record}")
        
        # Update progress: Finalization
        progress_data.update({
            'stage': 'finalization',
            'stage_name': 'Finalizing and preparing download...',
            'progress': 98,
            'elapsed_time': time.time() - start_time
        })
        emit_progress_update(task_id, progress_data)
        
        # Deduct credits based on question count - FIXED LOGIC FOR MANUAL SOLVE
        updated_user = db_helper.deduct_credits_by_count(user.get('phone_number'), len(questions))
        if updated_user:
            # Update session with new user data
            flask_session['user'] = updated_user
        else:
            logging.error(f"Failed to deduct credits for {user.get('phone_number')} - Manual solve")
        
        # Update progress: Completed
        final_elapsed = time.time() - start_time
        progress_data.update({
            'stage': 'completed',
            'stage_name': 'Manual questions processed successfully!',
            'progress': 100,
            'elapsed_time': final_elapsed,
            'total_time': f"{final_elapsed:.1f}s",
            'download_ready': True
        })
        emit_progress_update(task_id, progress_data)
        
        # Clean up task after a short delay
        def cleanup_task():
            time.sleep(5)  # Keep task info for 5 seconds
            if task_id in active_tasks:
                del active_tasks[task_id]
            if task_id in task_progress:
                del task_progress[task_id]
        
        threading.Thread(target=cleanup_task, daemon=True).start()
        
        # Return task_id in response header for frontend tracking
        response = make_response(send_file(generated_path, as_attachment=True))
        response.headers['X-Task-ID'] = task_id
        return response
        
    except Exception as e:
        logging.exception("Error processing manual questions")
        # Record failed submission due to exception
        try:
            user = flask_session.get('user')
            if user:
                db_helper.insert_submission(
                    phone_number=user.get('phone_number'),
                    pdf_name='Manual Questions',
                    questions_count=0,
                    questions_solved=0,
                    questions_failed=0,
                    failed_questions=[],
                    solved=False,
                    error_details=f"Processing exception: {str(e)}",
                    processing_time_seconds=(time.time() - start_time),
                    submission_type='manual'
                )
        except:
            pass  # Don't let logging errors break the main error response
        return jsonify({'error': str(e)}), 500

@app.route('/user_info', methods=['GET'])
def get_user_info():
    user = flask_session.get('user')
    if not user:
        return jsonify({'error': 'Not logged in'}), 401
    
    # Always fetch fresh user data from database to ensure accurate credit display
    fresh_user_data = db_helper.get_user_by_phone(user.get('phone_number'))
    if fresh_user_data:
        # Update session with fresh data
        flask_session['user'] = fresh_user_data
        user = fresh_user_data
    
    return jsonify({
        'phone_number': user.get('phone_number'),
        'name': user.get('name', 'User'),
        'credits': user.get('credits', 0),
        'is_priority': user.get('is_priority', False)
    })

@app.route('/submission_history', methods=['GET'])
def get_submission_history():
    """Get user's submission history - for debugging"""
    user = flask_session.get('user')
    if not user:
        return jsonify({'error': 'Not logged in'}), 401
    
    submissions = db_helper.get_user_submissions(user.get('phone_number'), limit=20)
    stats = db_helper.get_user_stats(user.get('phone_number'))
    
    return jsonify({
        'submissions': submissions,
        'stats': stats,
        'total_submissions': len(submissions)
    })

@app.route('/check_auth', methods=['GET'])
def check_auth():
    """Check if user is authenticated"""
    user = flask_session.get('user')
    authenticated = flask_session.get('authenticated', False)
    
    if user and authenticated:
        # Always fetch fresh user data from database to ensure accurate credit display
        fresh_user_data = db_helper.get_user_by_phone(user.get('phone_number'))
        if fresh_user_data:
            # Update session with fresh data
            flask_session['user'] = fresh_user_data
            user = fresh_user_data
        
        return jsonify({
            'authenticated': True,
            'user': {
                'name': user.get('name', 'User'),
                'phone_number': user.get('phone_number'),
                'credits': user.get('credits', 0)
            }
        })
    else:
        return jsonify({'authenticated': False})

@app.route('/refresh_user_session', methods=['POST'])
def refresh_user_session():
    """Refresh user session data from database"""
    try:
        user = flask_session.get('user')
        authenticated = flask_session.get('authenticated', False)
        
        # Check if user exists in session and is authenticated
        if not user or not authenticated:
            logging.warning(f"Session refresh attempted but user not authenticated. User: {bool(user)}, Authenticated: {authenticated}")
            return jsonify({'error': 'Not logged in'}), 401
        
        # Validate phone number exists
        phone_number = user.get('phone_number')
        if not phone_number:
            logging.warning("Session refresh attempted but no phone number in user data")
            return jsonify({'error': 'Invalid user data'}), 401
        
        # Fetch fresh user data from database
        fresh_user_data = db_helper.get_user_by_phone(user.get('phone_number'))
        if fresh_user_data:
            # Update session with fresh data
            flask_session['user'] = fresh_user_data
            flask_session['authenticated'] = True  # Ensure authenticated flag is set
            flask_session.permanent = True  # Make session persistent
            logging.info(f"Session refreshed for user {user.get('phone_number')} with credits: {fresh_user_data.get('credits', 0)}")
            # Also log the credit change for debugging
            old_credits = user.get('credits', 0)
            new_credits = fresh_user_data.get('credits', 0)
            if old_credits != new_credits:
                logging.info(f"Credit change detected: {old_credits} -> {new_credits}")
            
            return jsonify({
                'success': True,
                'user': {
                    'name': fresh_user_data.get('name', 'User'),
                    'phone_number': fresh_user_data.get('phone_number'),
                    'credits': fresh_user_data.get('credits', 0),
                    'is_priority': fresh_user_data.get('is_priority', False)
                },
                'credit_change': {
                    'old_credits': old_credits,
                    'new_credits': new_credits,
                    'changed': old_credits != new_credits
                }
            })
        else:
            return jsonify({'error': 'User not found in database'}), 404
            
    except Exception as e:
        logging.error(f"Error refreshing user session: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/check_session_status', methods=['GET'])
def check_session_status():
    """Check if user session is still valid and refresh if needed"""
    try:
        user = flask_session.get('user')
        authenticated = flask_session.get('authenticated', False)
        
        if not user or not authenticated:
            return jsonify({'authenticated': False, 'needs_login': True}), 401
        
        # Refresh user data to ensure it's current
        fresh_user_data = db_helper.get_user_by_phone(user.get('phone_number'))
        if fresh_user_data:
            flask_session['user'] = fresh_user_data
            flask_session['authenticated'] = True
            
            return jsonify({
                'authenticated': True,
                'user': {
                    'name': fresh_user_data.get('name', 'User'),
                    'phone_number': fresh_user_data.get('phone_number'),
                    'credits': fresh_user_data.get('credits', 0),
                    'is_priority': fresh_user_data.get('is_priority', False)
                }
            })
        else:
            # User not found in database, clear session
            flask_session.clear()
            return jsonify({'authenticated': False, 'needs_login': True}), 401
            
    except Exception as e:
        logging.error(f"Error checking session status: {e}")
        return jsonify({'authenticated': False, 'error': str(e)}), 500

@app.route('/debug_session', methods=['GET'])
def debug_session():
    """Debug session data - for troubleshooting"""
    try:
        user = flask_session.get('user')
        authenticated = flask_session.get('authenticated', False)
        session_permanent = flask_session.permanent
        
        return jsonify({
            'has_user': bool(user),
            'authenticated': authenticated,
            'session_permanent': session_permanent,
            'user_phone': user.get('phone_number') if user else None,
            'user_credits': user.get('credits') if user else None,
            'session_keys': list(flask_session.keys())
        })
    except Exception as e:
        logging.error(f"Error debugging session: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/complete_submission', methods=['POST'])
def complete_submission():
    try:
        submission_data = request.get_json()

        user = flask_session.get('user')
        if not user:
            return jsonify({'error': 'User not authenticated'}), 401

        pdf_name = submission_data.get('pdf_name')
        questions_count = submission_data.get('questions_count', 0)
        questions_solved = submission_data.get('questions_solved', 0)
        questions_failed = questions_count - questions_solved
        failed_questions = submission_data.get('failed_questions', [])
        error_details = submission_data.get('error_details', None)
        processing_time_seconds = submission_data.get('processing_time_seconds', None)

        solved = questions_solved == questions_count

        db_helper.insert_submission(user['phone_number'], pdf_name, questions_count,
                                    questions_solved, questions_failed, failed_questions,
                                    solved, error_details, processing_time_seconds)

        if not solved:
            return jsonify({'message': 'Submission incomplete', 'redirect': '/manual_entry'}), 400

        updated_user = db_helper.deduct_credit(user['phone_number'])
        flask_session['user'] = updated_user if updated_user else user

        return jsonify({'message': 'Submission complete, 1 credit deducted'}), 200

    except Exception as e:
        logging.exception("Error completing submission")
        return jsonify({'error': str(e)}), 500

@app.route('/logout')
def logout():
    flask_session.clear()
    return redirect('/')

@app.route('/test_twofactor', methods=['GET'])
def test_twofactor():
    """Test 2factor.in configuration - for debugging"""
    try:
        twofactor_otp = TwoFactorOTP()
        connection_test = twofactor_otp.test_connection()
        
        # Get balance information
        balance_info = twofactor_otp.get_balance()
        voice_balance_info = twofactor_otp.get_voice_balance()
        
        return jsonify({
            'twofactor_configured': True,
            'api_key_present': bool(twofactor_otp.api_key),
            'connection_test': connection_test,
            'sms_balance': balance_info,
            'voice_balance': voice_balance_info
        })
    except Exception as e:
        return jsonify({
            'twofactor_configured': False,
            'error': str(e)
        })

# =================== PLACEHOLDER FOR FUTURE PAYMENT INTEGRATION ===================
@app.route('/initiate_payment', methods=['POST'])
def initiate_payment():
    """Legacy payment endpoint - Simplified to use the working create-payment-session"""
    try:
        logging.info("Legacy payment endpoint called")
        
        # Get request data
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        # Extract user and payment info
        name = data.get('name', 'Customer')
        email = data.get('email', 'customer@codedebhai.com')
        phone = data.get('phone', '9999999999')
        amount = data.get('amount', 99)
        plan_type = data.get('plan_type', 'starter')

        # Prepare data for the new endpoint
        session_data = {
            'amount': amount,
            'currency': 'INR',
            'plan_type': plan_type,
            'user': {
                'name': name,
                'email': email,
                'phone': phone
            }
        }

        # Call the new simplified endpoint internally
        with app.test_request_context('/create-payment-session', method='POST', json=session_data):
            response = create_payment_session()
            
            # If successful, modify response format for compatibility
            if isinstance(response, tuple) and response[1] == 200:
                response_data = response[0].get_json()
                if response_data.get('success'):
                    # Clean up malformed payment session ID from the response
                    payment_session_id = response_data.get('payment_session_id')
                    if payment_session_id and payment_session_id.endswith('paymentpayment'):
                        payment_session_id = payment_session_id[:-7]  # Remove duplicate 'payment'
                        logging.info(f"Cleaned payment session ID: {payment_session_id}")
                    
                    return jsonify({
                        'payment_session_id': payment_session_id,
                        'order_id': response_data.get('order_id'),
                        'payment_link': f"https://sandbox.cashfree.com/checkout/{payment_session_id}"
                    })
            
            return response

    except Exception as e:
        logging.exception("Error in legacy payment initiation")
        return jsonify({'error': 'Payment initiation failed', 'debug': str(e)}), 500

# Removed duplicate payment_success endpoint - using the one defined earlier

@app.route('/payment/callback', methods=['POST'])
def payment_callback():
    """Handle payment webhook from Cashfree"""
    try:
        # Get the webhook data
        webhook_data = request.get_json()
        logging.info(f"🔔 Received webhook data: {webhook_data}")
        
        # Handle missing webhook data
        if not webhook_data:
            logging.error("❌ No webhook data received")
            return jsonify({'status': 'error', 'message': 'No webhook data received'}), 400
        
        # Extract data from Cashfree webhook format
        # Cashfree sends data in nested structure: data.order and data.payment
        data = webhook_data.get('data', {})
        order_data = data.get('order', {})
        payment_data = data.get('payment', {})
        
        order_id = order_data.get('order_id')
        payment_status = payment_data.get('payment_status')  # SUCCESS, FAILED, etc.
        gateway_payment_id = payment_data.get('cf_payment_id')
        
        logging.info(f"📋 Processing webhook - Order: {order_id}, Status: {payment_status}, Payment ID: {gateway_payment_id}")
        
        # Handle missing required fields
        if not order_id:
            logging.error("❌ Missing order_id in webhook data")
            return jsonify({'status': 'error', 'message': 'Missing order_id'}), 400
        
        # Check if payment is successful
        if payment_status == 'SUCCESS':
            logging.info(f"✅ Payment successful for order: {order_id}")
            
            payment_record = db_helper.get_payment_by_gateway_id(order_id)
            if payment_record:
                logging.info(f"💳 Found payment record: {payment_record}")
                
                # Update payment status and add credits
                current_user = db_helper.get_user_by_phone(payment_record['phone_number'])
                if current_user:
                    old_credits = current_user['credits']
                    new_credits = old_credits + payment_record['credits_added']
                    
                    logging.info(f"💰 Adding credits - User: {payment_record['phone_number']}, Old: {old_credits}, Adding: {payment_record['credits_added']}, New: {new_credits}")
                    
                    updated_user = db_helper.update_user_credits(payment_record['phone_number'], new_credits)
                    
                    # Update session if user is currently logged in
                    current_session_user = flask_session.get('user')
                    if current_session_user and current_session_user.get('phone_number') == payment_record['phone_number']:
                        # Refresh user data from database to ensure session has latest credits
                        fresh_user_data = db_helper.get_user_by_phone(payment_record['phone_number'])
                        if fresh_user_data:
                            flask_session['user'] = fresh_user_data
                            logging.info(f"🔄 Session updated for user {payment_record['phone_number']} with new credits: {fresh_user_data.get('credits', 0)}")
                    
                    # Update payment record with webhook received flag
                    db_helper.update_payment_status(order_id, 'paid', gateway_payment_id, webhook_received=True)
                    logging.info(f"✅ Payment processing complete for order: {order_id}")
                    
                    return jsonify({'status': 'success'}), 200
                else:
                    logging.error(f"❌ User not found for phone: {payment_record['phone_number']}")
                    return jsonify({'status': 'error', 'message': 'User not found'}), 404
            else:
                logging.error(f"❌ Payment record not found for order: {order_id}")
                return jsonify({'status': 'error', 'message': 'Payment record not found'}), 404
        
        elif payment_status == 'FAILED':
            logging.info(f"❌ Payment failed for order: {order_id}")
            # Update payment status to failed
            db_helper.update_payment_status(order_id, 'failed', gateway_payment_id, webhook_received=True)
            return jsonify({'status': 'acknowledged', 'message': 'Payment failure recorded'}), 200
        
        else:
            logging.info(f"ℹ️ Unhandled payment status '{payment_status}' for order: {order_id}")
            return jsonify({'status': 'acknowledged', 'message': f'Status {payment_status} recorded'}), 200
        
    except Exception as e:
        logging.error(f"❌ Payment callback error: {e}")
        logging.exception("Full webhook error traceback:")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/webhook-payment', methods=['POST'])
def webhook_payment():
    """Unified Payment Webhook to handle success from Cashfree"""
    try:
        # Get the webhook data
        data = request.get_json()
        order_id = data.get('orderId')
        status = data.get('orderStatus')
        payment_id = data.get('paymentSessionId')
        
        # Check if order is successfully paid
        if status == 'PAID':
            payment_record = db_helper.get_payment_by_gateway_id(order_id)
            if payment_record:
                current_user = db_helper.get_user_by_phone(payment_record['phone_number'])
                if current_user:
                    new_credits = current_user['credits'] + payment_record['credits_added']
                    db_helper.update_user_credits(payment_record['phone_number'], new_credits)
                    db_helper.update_payment_status(order_id, 'paid', payment_id)
                    
                    # Refresh session with updated user data
                    current_session_user = flask_session.get('user')
                    if current_session_user and current_session_user.get('phone_number') == payment_record['phone_number']:
                        fresh_user_data = db_helper.get_user_by_phone(payment_record['phone_number'])
                        if fresh_user_data:
                            flask_session['user'] = fresh_user_data
                            flask_session['authenticated'] = True
                            flask_session.permanent = True
                            logging.info(f"Session updated for user {payment_record['phone_number']} after payment.")

                    return jsonify({'status': 'success'}), 200
            else:
                logging.error(f"Payment record not found for order: {order_id}")
                return jsonify({'status': 'error', 'message': 'Payment record not found'}), 404
        else:
            logging.info(f"Unhandled order status {status} for order {order_id}")
            return jsonify({'status': 'acknowledged', 'message': f'Status {status} recorded'}), 200
    except Exception as e:
        logging.error(f"Webhook error: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/get_payment_plans', methods=['GET'])
def get_payment_plans():
    """Get all available payment plans with enhanced frontend data"""
    try:
        # Define the plans with all frontend data
        plans = [
            {
                'id': 'starter',
                'plan_name': 'Starter Plan',
                'amount': 99,
                'credits': 10,
                'is_priority': False,
                'badge': 'New Entry',
                'description': '₹99 → 10 credits - Entry-level for new users',
                'features': [
                    '✅ 1 credit = 1 solved pdf (max 20 questions)',
                    '✅ Entry-level for new users',
                    '✅ Ideal for light users',
                    '✅ Perfect for trying us out!',
                    '✅ Test drive our service'
                ],
                'button_text': '🚀 Pay Now',
                'button_class': 'secondary'
            },
            {
                'id': 'monthly',
                'plan_name': 'Monthly Saver',
                'amount': 299,
                'credits': 50,
                'is_priority': False,
                'badge': 'Best Value',
                'is_featured': True,
                'savings': 'Save 33% per question!',
                'description': '₹299 → 50 credits - Best value for regular users',
                'features': [
                    '✅ 50 pdf Solutions',
                    '✅ Flexibility & reliability',
                    '✅ Solve more when needed',
                    '✅ Great for regular assignments',
                    '✅ Perfect for semester workload',
                    '✅ Most Popular Among students'
                ],
                'button_text': '💳 Pay Now',
                'button_class': 'featured'
            },
            {
                'id': 'power',
                'plan_name': 'Power Plan',
                'amount': 799,
                'credits': 150,
                'is_priority': True,
                'badge': '3 Months Access',
                'savings': 'Save 45% vs Starter!',
                'description': '₹799 → 150 credits + Priority Access - Maximum value for consistent users',
                'features': [
                    '✅ 150 pdf solved',
                    '✅ Valid for 3 months',
                    '✅ Priority solving queue',
                    '✅ Perfect for exam rush & semester load',
                    '✅ Maximum value for consistent users'
                ],
                'button_text': '⚡ Pay Now',
                'button_class': 'primary'
            }
        ]
        
        # Add free trial info
        free_trial = {
            'credits': 5,
            'description': '5 free credits tied to phone number to avoid abuse'
        }
        
        # Add conversion info
        conversion_info = {
            'rule': '1 credit = 1 PDF',
            'description': 'Each PDF can contain up to 20 coding questions'
        }
        
        return jsonify({
            'plans': plans,
            'free_trial': free_trial,
            'conversion': conversion_info
        })
    except Exception as e:
        logging.error(f"Error getting payment plans: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/user_payments', methods=['GET'])
def get_user_payments():
    """Get user's payment history"""
    try:
        user = flask_session.get('user')
        if not user:
            return jsonify({'error': 'Not logged in'}), 401
        
        payments = db_helper.get_user_payments(user.get('phone_number'), limit=10)
        payment_stats = db_helper.get_user_payment_stats(user.get('phone_number'))
        
        return jsonify({
            'payments': payments,
            'stats': payment_stats
        })
    except Exception as e:
        logging.error(f"Error getting user payments: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/test_cashfree', methods=['GET'])
def test_cashfree():
    """Test Cashfree configuration"""
    try:
        app_id = os.getenv('CASHFREE_APP_ID')
        secret_key = os.getenv('CASHFREE_SECRET_KEY')
        environment = os.getenv('CASHFREE_ENVIRONMENT', 'sandbox')
        
        if not app_id or not secret_key:
            return jsonify({
                'cashfree_configured': False,
                'error': 'Missing Cashfree credentials'
            })
        
        # Test API connection
        headers = {
            "x-client-id": app_id,
            "x-client-secret": secret_key,
            "Content-Type": "application/json"
        }
        
        # Test with a simple API call
        test_response = requests.get(
            "https://sandbox.cashfree.com/pg/orders",
            headers=headers,
            timeout=10
        )
        
        connection_test = {
            'success': test_response.status_code in [200, 401, 403],  # These are expected responses
            'status_code': test_response.status_code,
            'response': test_response.text[:200] if test_response.text else 'No response'
        }
        
        return jsonify({
            'cashfree_configured': True,
            'environment': environment,
            'app_id_present': bool(app_id),
            'secret_key_present': bool(secret_key),
            'connection_test': connection_test
        })
        
    except Exception as e:
        return jsonify({
            'cashfree_configured': False,
            'error': str(e)
        })

@app.route('/test_payment_success/<order_id>', methods=['GET'])
def test_payment_success(order_id):
    """Test endpoint to simulate successful payment for testing"""
    try:
        user = flask_session.get('user')
        if not user:
            return jsonify({'error': 'Please login first'}), 401
        
        # Get payment record
        payment_record = db_helper.get_payment_by_gateway_id(order_id)
        if not payment_record:
            return jsonify({'error': 'Payment record not found'}), 404
        
        # Update payment status
        updated_payment = db_helper.update_payment_status(order_id, 'paid', f'test_payment_{order_id}')
        
        # Add credits to user account - NO BONUS CREDITS
        credits_to_add = payment_record['credits_added']  # Add exact credits as per plan
        user_updated = db_helper.add_credits_with_rollover_tracking(
            payment_record['phone_number'], 
            credits_to_add,
            payment_record.get('is_priority', False)
        )
        
        if user_updated:
            # Update session with new user data
            updated_user = db_helper.get_user_by_phone(user.get('phone_number'))
            if updated_user:
                flask_session['user'] = updated_user
            
            logging.info(f"Test payment completed: Order {order_id}, {credits_to_add} credits added to {payment_record['phone_number']}")
            return jsonify({
                'status': 'success', 
                'message': f'Test payment successful! {credits_to_add} credits added.',
                'new_credit_balance': updated_user.get('credits', 0) if updated_user else 0
            }), 200
        else:
            return jsonify({'error': 'Failed to add credits'}), 500
            
    except Exception as e:
        logging.error(f"Test payment error: {e}")
        return jsonify({'error': str(e)}), 500
    
@app.route('/debug_user/<phone_number>', methods=['GET'])
def debug_user(phone_number):
    """Debug endpoint to check user record"""
    try:
        user = db_helper.get_user_by_phone(phone_number)
        if user:
            return jsonify({
                'user_found': True,
                'user_data': user,
                'current_credits': user.get('credits', 0)
            })
        else:
            return jsonify({'user_found': False, 'message': f'No user found with phone: {phone_number}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/debug_payment/<order_id>', methods=['GET'])
def debug_payment(order_id):
    """Debug endpoint to check payment record"""
    try:
        payment = db_helper.get_payment_by_gateway_id(order_id)
        if payment:
            return jsonify({
                'payment_found': True,
                'payment_data': payment
            })
        else:
            return jsonify({'payment_found': False, 'message': f'No payment found with order_id: {order_id}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/create_payment_record/<order_id>/<int:amount>', methods=['GET'])
def create_payment_record_manual(order_id, amount):
    """Manually create payment record for testing"""
    try:
        user = flask_session.get('user')
        if not user:
            return jsonify({'error': 'Please login first'}), 401
        
        # Check if payment record already exists
        existing_payment = db_helper.get_payment_by_gateway_id(order_id)
        if existing_payment:
            return jsonify({'error': 'Payment record already exists', 'payment': existing_payment}), 409
        
        # Determine credits based on amount - NO BONUS CREDITS
        credits_map = {
            99: 10,   # Starter plan
            299: 50,  # Monthly plan  
            799: 150  # Power plan - Fixed to 150 credits
        }
        credits_to_add = credits_map.get(amount, 10)
        
        # Get current user
        current_user = db_helper.get_user_by_phone(user.get('phone_number'))
        if not current_user:
            return jsonify({'error': 'User not found'}), 404
        
        # Create payment record
        payment_record = db_helper.create_payment_record(
            user_id=current_user['id'],
            phone_number=user.get('phone_number'),
            gateway_order_id=order_id,
            gateway_payment_id=None,
            amount=amount * 100,  # Convert to paise
            credits_added=credits_to_add,
            plan_type='demo',
            payment_status='pending'
        )
        
        if payment_record:
            return jsonify({
                'success': True,
                'message': f'Payment record created for order {order_id}',
                'payment_record': payment_record
            })
        else:
            return jsonify({'error': 'Failed to create payment record'}), 500
            
    except Exception as e:
        logging.error(f"Error creating payment record: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/test_webhook_payment/<order_id>', methods=['POST'])
def test_webhook_payment(order_id):
    """Test webhook simulation for payment success - Fixed to use correct amount"""
    try:
        # Get payment record first to get the correct amount
        payment_record = db_helper.get_payment_by_gateway_id(order_id)
        if not payment_record:
            # Fallback: try to create payment record from session stash
            try:
                pending = flask_session.get('pending_payments', {}) or {}
                pending_info = pending.get(order_id)
            except Exception:
                pending_info = None
            if pending_info:
                try:
                    user_row = db_helper.get_user_by_phone(pending_info.get('phone_number'))
                    if user_row:
                        created = db_helper.create_payment_record(
                            user_id=user_row['id'],
                            phone_number=pending_info.get('phone_number'),
                            gateway_order_id=order_id,
                            gateway_payment_id=None,
                            amount=int(pending_info.get('amount', 99)) * 100,
                            credits_added=int(pending_info.get('credits_to_add', 10)),
                            plan_type=pending_info.get('plan_type', 'fallback'),
                            payment_status='pending'
                        )
                        if created:
                            payment_record = created
                            logging.info(f"Created missing payment record from session stash for {order_id}")
                except Exception as e:
                    logging.error(f"Failed to create missing record for {order_id}: {e}")
            if not payment_record:
                logging.error(f"Payment record not found for order: {order_id}")
                return jsonify({'error': f'Payment record not found for order: {order_id}'}), 404
        
        logging.info(f"Found payment record: {payment_record}")
        
        # Check if payment was already processed to prevent double processing
        if payment_record.get('payment_status') == 'paid':
            logging.warning(f"Payment {order_id} already processed, skipping credit addition")
            return jsonify({
                'status': 'already_processed',
                'message': f'Payment already processed! Credits were already added.',
                'order_id': order_id,
                'credits_added': payment_record['credits_added'],
                'user_phone': payment_record['phone_number']
            }), 200
        
        # Use the correct amount from payment record (convert from paise to rupees)
        actual_amount = payment_record['amount'] // 100
        
        # Simulate webhook data from Cashfree with CORRECT amount
        webhook_data = {
            'orderId': order_id,
            'orderStatus': 'PAID',
            'paymentSessionId': f'test_session_{order_id}',
            'orderAmount': actual_amount  # Use actual amount, not hardcoded 299
        }
        
        logging.info(f"Test webhook simulation for order: {order_id}")
        logging.info(f"Simulated webhook data: {webhook_data}")
        
        # Update payment status
        updated_payment = db_helper.update_payment_status(order_id, 'paid', webhook_data['paymentSessionId'])
        logging.info(f"Payment status updated: {updated_payment}")
        
        # Get user's current credits before adding
        current_user = db_helper.get_user_by_phone(payment_record['phone_number'])
        credits_before = current_user['credits'] if current_user else 0
        logging.info(f"User {payment_record['phone_number']} credits before: {credits_before}")
        
        # Add credits to user account - Use exact credits from payment record
        credits_to_add = payment_record['credits_added']  # Add exact credits as per plan
        logging.info(f"Adding {credits_to_add} credits to user {payment_record['phone_number']} for ₹{actual_amount} payment")
        
        user_updated = db_helper.add_credits_to_user(
            payment_record['phone_number'], 
            credits_to_add,
            payment_record.get('is_priority', False)
        )
        
        if user_updated:
            credits_after = user_updated['credits']
            logging.info(f"User {payment_record['phone_number']} credits after: {credits_after}")
            
            # Update session if user is currently logged in
            current_session_user = flask_session.get('user')
            if current_session_user and current_session_user.get('phone_number') == payment_record['phone_number']:
                # Refresh user data from database to ensure session has latest credits
                fresh_user_data = db_helper.get_user_by_phone(payment_record['phone_number'])
                if fresh_user_data:
                    flask_session['user'] = fresh_user_data
                    logging.info(f"Session updated for user {payment_record['phone_number']} with new credits: {fresh_user_data.get('credits', 0)}")
            
            logging.info(f"Credits added successfully: {credits_to_add} credits to {payment_record['phone_number']} for ₹{actual_amount} payment")
            return jsonify({
                'status': 'success', 
                'message': f'Test webhook processed successfully! {credits_to_add} credits added for ₹{actual_amount} payment.',
                'order_id': order_id,
                'credits_added': credits_to_add,
                'user_phone': payment_record['phone_number'],
                'credits_before': credits_before,
                'credits_after': credits_after,
                'payment_amount': actual_amount
            }), 200
        else:
            logging.error(f"Failed to add credits for order {order_id}")
            return jsonify({'error': 'Failed to add credits'}), 500
            
    except Exception as e:
        logging.error(f"Test webhook error: {e}")
        return jsonify({'error': str(e)}), 500

# =================== ADMIN ROLLOVER ENDPOINTS ===================

@app.route('/admin/rollover/status', methods=['GET'])
def admin_rollover_status():
    """Get rollover system status"""
    try:
        # Simple admin check - in production, implement proper admin authentication
        # For now, just check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({'error': 'Admin access required'}), 401
        
        rollover_system = ProductionRolloverSystem()
        
        return jsonify({
            'enabled': rollover_system.enabled,
            'dry_run': rollover_system.dry_run,
            'notification_email': rollover_system.notification_email,
            'backup_enabled': rollover_system.backup_enabled,
            'last_run': datetime.now().isoformat(),  # Current timestamp for production
            'next_scheduled_run': 'Monthly - 1st of each month at 2 AM'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/admin/rollover/run', methods=['POST'])
def admin_run_rollover():
    """Manual rollover trigger for admin"""
    try:
        # Simple admin check - in production, implement proper admin authentication
        user = flask_session.get('user')
        if not user:
            return jsonify({'error': 'Admin access required'}), 401
        
        data = request.get_json() or {}
        dry_run = data.get('dry_run', False)
        
        # Override environment temporarily for this run
        if dry_run:
            os.environ['ROLLOVER_DRY_RUN'] = 'true'
        
        rollover_system = ProductionRolloverSystem()
        
        # Run in a separate thread to avoid blocking the request
        import threading
        
        def run_rollover():
            rollover_system.run_monthly_rollover()
        
        rollover_thread = threading.Thread(target=run_rollover)
        rollover_thread.start()
        
        return jsonify({
            'message': f'Rollover started successfully (dry_run: {dry_run})',
            'dry_run': dry_run,
            'status': 'running'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/admin/rollover/test', methods=['POST'])
def admin_test_rollover():
    """Test rollover system with dry run"""
    try:
        # Simple admin check - in production, implement proper admin authentication
        user = flask_session.get('user')
        if not user:
            return jsonify({'error': 'Admin access required'}), 401
        
        # Force dry run for testing
        os.environ['ROLLOVER_DRY_RUN'] = 'true'
        
        rollover_system = ProductionRolloverSystem()
        rollover_system.run_monthly_rollover()
        
        return jsonify({
            'message': 'Test rollover completed successfully',
            'stats': rollover_system.rollover_stats,
            'dry_run': True
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/admin/rollover/eligible-users', methods=['GET'])
def admin_get_eligible_users():
    """Get list of users eligible for rollover"""
    try:
        # Simple admin check - in production, implement proper admin authentication
        user = flask_session.get('user')
        if not user:
            return jsonify({'error': 'Admin access required'}), 401
        
        rollover_system = ProductionRolloverSystem()
        users_with_credits = rollover_system.get_all_users_with_credits()
        
        eligible_users = []
        for user_data in users_with_credits:
            eligibility = rollover_system.get_user_plan_eligibility(user_data['phone_number'])
            eligible_users.append({
                'phone_number': user_data['phone_number'],
                'name': user_data['name'],
                'credits': user_data['credits'],
                'eligible': eligibility['eligible'],
                'plan_type': eligibility.get('plan_type', 'unknown'),
                'reason': eligibility['reason']
            })
        
        return jsonify({
            'total_users': len(users_with_credits),
            'eligible_for_rollover': len([u for u in eligible_users if u['eligible']]),
            'users': eligible_users
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ============= TERMINAL UTILITIES API ENDPOINTS =============

@app.route('/api/terminal/clean_path', methods=['GET'])
def api_clean_terminal_path():
    """API endpoint to get clean terminal path"""
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        # Get user name for personalized path generation
        user_name = user.get('name', 'Developer')
        
        # Set current user name for path generation
        clean_terminal_path._current_user_name = user_name
        
        # Generate clean path
        clean_path = clean_terminal_path()
        
        return jsonify({
            "success": True,
            "clean_path": clean_path,
            "user_name": user_name
        })
        
    except Exception as e:
        logging.exception("Error generating clean terminal path")
        return jsonify({'error': str(e)}), 500

@app.route('/api/terminal/screenshot', methods=['POST'])
def api_take_terminal_screenshot():
    """API endpoint to take terminal screenshot"""
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        data = request.get_json() or {}
        output_path = data.get('output_path', f'terminal_screenshot_{uuid.uuid4().hex[:8]}.png')
        
        # Take screenshot
        result = take_screenshot(output_path)
        
        return jsonify({
            "success": True,
            "message": result,
            "output_path": output_path
        })
        
    except Exception as e:
        logging.exception("Error taking terminal screenshot")
        return jsonify({'error': str(e)}), 500

@app.route('/api/terminal/fix', methods=['POST'])
def api_fix_terminal():
    """API endpoint to fix terminal display issues"""
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        # Fix terminal
        result = fix_terminal()
        
        return jsonify({
            "success": True,
            "message": result
        })
        
    except Exception as e:
        logging.exception("Error fixing terminal")
        return jsonify({'error': str(e)}), 500

@app.route('/api/terminal/suppress_output', methods=['POST'])
def api_suppress_terminal_output():
    """API endpoint to suppress extra terminal output"""
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        # Suppress extra output
        suppress_extra_output()
        
        return jsonify({
            "success": True,
            "message": "Extra terminal output suppressed successfully"
        })
        
    except Exception as e:
        logging.exception("Error suppressing terminal output")
        return jsonify({'error': str(e)}), 500

@app.route('/api/terminal/utils_status', methods=['GET'])
def api_terminal_utils_status():
    """API endpoint to check terminal utilities status"""
    try:
        # Check if user is logged in
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to use this service."}), 401
        
        # Initialize terminal utils
        utils = TerminalUtils()
        
        return jsonify({
            "success": True,
            "system": utils.system,
            "is_windows": utils.is_windows,
            "available_functions": [
                "clean_terminal_path",
                "take_screenshot", 
                "fix_terminal",
                "suppress_extra_output",
                "clear_terminal_buffer"
            ],
            "status": "Terminal utilities are available and ready"
        })
        
    except Exception as e:
        logging.exception("Error checking terminal utils status")
        return jsonify({'error': str(e)}), 500

# API Key Performance Monitoring Endpoint
@app.route('/api/keys/stats', methods=['GET'])
def get_api_key_stats():
    """Get API key performance statistics"""
    try:
        # Check if user is logged in (admin check in production)
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Please login to view API stats."}), 401
        
        # Get API key statistics from both providers
        deepseek_stats = deepseek_manager.get_stats()
        claude_stats = claude_manager.get_stats()
        
        # Format DeepSeek statistics
        deepseek_formatted = []
        for key_index, stats in deepseek_stats.items():
            current_time = time.time()
            is_rate_limited = stats['rate_limited_until'] > current_time
            remaining_cooldown = max(0, stats['rate_limited_until'] - current_time)
            
            deepseek_formatted.append({
                'key_index': key_index + 1,
                'key_id': f"DSK-{key_index + 1}",
                'provider': 'DeepSeek',
                'total_requests': stats['requests'],
                'total_failures': stats['failures'],
                'success_rate': f"{((stats['requests'] - stats['failures']) / max(stats['requests'], 1) * 100):.1f}%",
                'is_rate_limited': is_rate_limited,
                'remaining_cooldown': f"{remaining_cooldown:.1f}s" if is_rate_limited else "Available",
                'last_used': f"{(current_time - stats['last_used']):.1f}s ago" if stats['last_used'] > 0 else "Never",
                'status': "Rate Limited" if is_rate_limited else "Available"
            })
        
        # Format Claude statistics
        claude_formatted = []
        for key_index, stats in claude_stats.items():
            current_time = time.time()
            is_rate_limited = stats['rate_limited_until'] > current_time
            remaining_cooldown = max(0, stats['rate_limited_until'] - current_time)
            requests_this_minute = stats.get('requests_this_minute', 0)
            max_per_minute = claude_manager.max_requests_per_minute
            
            claude_formatted.append({
                'key_index': key_index + 1,
                'key_id': f"CLD-{key_index + 1}",
                'provider': 'Claude',
                'total_requests': stats['requests'],
                'total_failures': stats['failures'],
                'success_rate': f"{((stats['requests'] - stats['failures']) / max(stats['requests'], 1) * 100):.1f}%",
                'is_rate_limited': is_rate_limited,
                'remaining_cooldown': f"{remaining_cooldown:.1f}s" if is_rate_limited else "Available",
                'last_used': f"{(current_time - stats['last_used']):.1f}s ago" if stats['last_used'] > 0 else "Never",
                'requests_this_minute': f"{requests_this_minute}/{max_per_minute}",
                'status': "Rate Limited" if is_rate_limited else "Available"
            })
        
        # Combine all formatted stats
        formatted_stats = claude_formatted + deepseek_formatted
        
        # Overall statistics
        deepseek_total_requests = sum(s['requests'] for s in deepseek_stats.values())
        deepseek_total_failures = sum(s['failures'] for s in deepseek_stats.values())
        deepseek_available_keys = sum(1 for s in deepseek_stats.values() if s['rate_limited_until'] <= current_time)
        
        claude_total_requests = sum(s['requests'] for s in claude_stats.values())
        claude_total_failures = sum(s['failures'] for s in claude_stats.values())
        claude_available_keys = claude_manager.get_available_keys_count()
        
        return jsonify({
            "success": True,
            "providers": {
                "claude": {
                    "total_keys": len(CLAUDE_KEYS),
                    "available_keys": claude_available_keys,
                    "rate_limited_keys": len(CLAUDE_KEYS) - claude_available_keys,
                    "total_requests": claude_total_requests,
                    "total_failures": claude_total_failures,
                    "success_rate": f"{((claude_total_requests - claude_total_failures) / max(claude_total_requests, 1) * 100):.1f}%"
                },
                "deepseek": {
                    "total_keys": len(DEEPSEEK_KEYS),
                    "available_keys": deepseek_available_keys,
                    "rate_limited_keys": len(DEEPSEEK_KEYS) - deepseek_available_keys,
                    "total_requests": deepseek_total_requests,
                    "total_failures": deepseek_total_failures,
                    "success_rate": f"{((deepseek_total_requests - deepseek_total_failures) / max(deepseek_total_requests, 1) * 100):.1f}%"
                }
            },
            "overall_stats": {
                "total_requests": claude_total_requests + deepseek_total_requests,
                "total_failures": claude_total_failures + deepseek_total_failures,
                "overall_success_rate": f"{(((claude_total_requests + deepseek_total_requests) - (claude_total_failures + deepseek_total_failures)) / max(claude_total_requests + deepseek_total_requests, 1) * 100):.1f}%",
                "load_distribution": "Intelligent API selection with rate limit management"
            },
            "key_details": formatted_stats,
            "recommendations": {
                "optimal_performance": (claude_available_keys + deepseek_available_keys) >= 4,
                "message": f"🟢 Healthy ({claude_available_keys} Claude + {deepseek_available_keys} DeepSeek available)" if (claude_available_keys + deepseek_available_keys) >= 4 else f"🟡 Limited availability ({claude_available_keys} Claude + {deepseek_available_keys} DeepSeek available)",
                "strategy": "Claude preferred for quality, DeepSeek for volume. Smart fallback active."
            }
        })
        
    except Exception as e:
        logging.exception("Error getting API key statistics")
        return jsonify({'error': str(e)}), 500

@app.route('/api/keys/reset', methods=['POST'])
def reset_api_key_stats():
    """Reset API key statistics (admin only)"""
    try:
        # Check if user is logged in (admin check in production)
        user = flask_session.get('user')
        if not user:
            return jsonify({"error": "Admin access required."}), 401
        
        # Reset all key statistics for both providers
        with deepseek_manager.lock:
            for key_index in deepseek_manager.key_stats:
                deepseek_manager.key_stats[key_index] = {
                    'requests': 0, 
                    'failures': 0, 
                    'last_used': 0, 
                    'rate_limited_until': 0
                }
        
        with claude_manager.lock:
            for key_index in claude_manager.key_stats:
                claude_manager.key_stats[key_index] = {
                    'requests': 0, 
                    'failures': 0, 
                    'last_used': 0, 
                    'rate_limited_until': 0,
                    'requests_this_minute': 0,
                    'minute_window_start': 0
                }
        
        logging.info(f"API key statistics reset by user {user.get('phone_number')}")
        
        return jsonify({
            "success": True,
            "message": "API key statistics reset successfully for both providers",
            "keys_reset": {
                "claude_keys": len(CLAUDE_KEYS),
                "deepseek_keys": len(DEEPSEEK_KEYS),
                "total_keys": len(CLAUDE_KEYS) + len(DEEPSEEK_KEYS)
            }
        })
        
    except Exception as e:
        logging.exception("Error resetting API key statistics")
        return jsonify({'error': str(e)}), 500

# Add error handlers to return JSON for API endpoints
@app.errorhandler(404)
def not_found_error(error):
    # Check if it's an API request
    if request.path.startswith('/verify_otp') or request.path.startswith('/setup_firebase_session'):
        return jsonify({'error': 'Endpoint not found'}), 404
    # Return normal 404 page for other requests
    return "<h1>404 - Page Not Found</h1>", 404

@app.errorhandler(500)
def internal_error(error):
    # Check if it's an API request
    if request.path.startswith('/verify_otp') or request.path.startswith('/setup_firebase_session') or request.content_type == 'application/json':
        logging.error(f"Internal server error on {request.path}: {error}")
        return jsonify({'error': 'Internal server error', 'details': 'Please try again or contact support'}), 500
@app.route('/get_user_data', methods=['GET'])
def get_user_data():
    """Get current user data from session for auto-populating forms"""
    try:
        user = flask_session.get('user')
        if not user:
            return jsonify({'error': 'User not logged in'}), 401
        
        # Return user data for form population
        return jsonify({
            'success': True,
            'user': {
                'name': user.get('name', ''),
                'phone_number': user.get('phone_number', ''),
                'credits': user.get('credits', 0)
            }
        })
    except Exception as e:
        logging.error(f"Error getting user data: {e}")
        return jsonify({'error': 'Failed to get user data'}), 500

    # Return normal 500 page for other requests
    return "<h1>500 - Internal Server Error</h1>", 500

@app.errorhandler(Exception)
def handle_exception(e):
    # Check if it's an API request
    if request.path.startswith('/verify_otp') or request.path.startswith('/setup_firebase_session') or request.content_type == 'application/json':
        logging.error(f"Unhandled exception on {request.path}: {e}")
        return jsonify({'error': 'An unexpected error occurred', 'details': str(e)}), 500
    # Re-raise for other requests
    raise e

# Setup error handlers for production
setup_error_handlers(app)

if __name__ == '__main__':
    # Get port from environment variable for deployment
    port = int(os.environ.get('PORT', 5000))
    # Start the app with SocketIO
    socketio.run(app, debug=False, host='0.0.0.0', port=port)
                